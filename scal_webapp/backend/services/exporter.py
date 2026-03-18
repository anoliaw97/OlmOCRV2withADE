from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document


def export_json(path: Path, payload: dict):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)


def export_excel(path: Path, tables: list[dict], ml_df: pd.DataFrame):
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame(tables).to_excel(writer, sheet_name="tables_json", index=False)
        ml_df.to_excel(writer, sheet_name="ml_ready", index=False)


def export_word(path: Path, tables: list[dict], graph_image: str | None = None):
    doc = Document()
    doc.add_heading("SCAL Extraction Report", level=1)
    doc.add_paragraph("Generated from extracted JSON tables (offline pipeline).")

    if graph_image and Path(graph_image).exists():
        doc.add_paragraph("Graph image:")
        doc.add_picture(graph_image)

    for t in tables[:30]:
        doc.add_heading(f"{t.get('table_id')} | {t.get('extraction_type')}", level=2)
        doc.add_paragraph(f"File: {t.get('file_name')} | Page: {t.get('page_number')}")
        doc.add_paragraph(f"Title: {t.get('table_title')}")

    doc.save(path)
