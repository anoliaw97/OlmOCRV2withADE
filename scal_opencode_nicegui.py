from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

import joblib
import pandas as pd
from nicegui import ui
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel

try:
    from bs4 import BeautifulSoup

    BS4_OK = True
except Exception:
    BS4_OK = False


# ---------------------------- CONFIG ----------------------------

DEFAULT_DATA_ROOT = Path(
    r"C:\Users\Mining\Downloads\Fine Tunining Datasets-20260318T052420Z-1-001\Fine Tunining Datasets\train"
)
EXPORT_DIR = Path("scal_local_exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)
INDEX_DIR = Path("scal_local_index")
INDEX_DIR.mkdir(parents=True, exist_ok=True)
PROMPT_FILE = Path("scal_saved_prompts.json")

DEFAULT_CHAT_PROMPTS = [
    {
        "name": "Extract certain columns only",
        "text": """You are a document analysis assistant.
Use only retrieved extracted evidence.
Return only Sample ID, Porosity, Depth in valid JSON.
Keep row order. Missing fields must be NULL.
No explanation, no extra text.""",
    },
    {
        "name": "Extract table based on keyword",
        "text": """Scan retrieved evidence and find the table containing keyword [keyword].
Return full table as JSON array with exact column names and row order.
If not found return: {"no_table": true}""",
    },
    {
        "name": "Graph extraction",
        "text": """From retrieved graph/table evidence, output structured numeric JSON.
No commentary. Numeric values only.""",
    },
]


# ---------------------------- STATE ----------------------------

class AppState:
    def __init__(self):
        self.data_root: Path = DEFAULT_DATA_ROOT
        self.docs: dict[str, dict[int, dict[str, Path]]] = {}
        self.current_doc: str | None = None

        self.index_namespace: str | None = None
        self.vectorizer: TfidfVectorizer | None = None
        self.matrix = None
        self.index_texts: list[str] = []
        self.index_meta: list[dict[str, Any]] = []

        self.chat_history: dict[str, list[tuple[str, str]]] = {"default": []}
        self.current_session: str = "default"

        self.last_hits: list[dict[str, Any]] = []
        self.last_render_tables: list[dict[str, Any]] = []

        self.llm_loaded = False
        self.vlm_loaded = False
        self.llm_model = "Qwen/Qwen2.5-3B-Instruct"

        self.vlm = None
        self.llm_tokenizer = None
        self.llm_model_obj = None


S = AppState()


# ---------------------------- HELPERS ----------------------------

def ts() -> str:
    return datetime.now().strftime("%H:%M:%S")


def parse_name(file_name: str) -> tuple[str | None, int | None, str]:
    m = re.match(r"^(.*)_page(\d+)\.(pdf|md|json)$", file_name, flags=re.IGNORECASE)
    if not m:
        return None, None, Path(file_name).suffix.lower().lstrip(".")
    return m.group(1), int(m.group(2)), m.group(3).lower()


def scan_dataset(root: Path) -> dict[str, dict[int, dict[str, Path]]]:
    docs: dict[str, dict[int, dict[str, Path]]] = {}
    if not root.exists():
        return docs
    for p in root.glob("*"):
        if not p.is_file():
            continue
        stem, page, ext = parse_name(p.name)
        if stem is None or page is None:
            continue
        docs.setdefault(stem, {}).setdefault(page, {})[ext] = p
    return docs


def coverage_info(pages: dict[int, dict[str, Path]]) -> dict[str, Any]:
    if not pages:
        return {"pdf_pages": 0, "extracted_pages": 0, "missing": []}
    all_pages = sorted(pages.keys())
    pdf_pages = [p for p in all_pages if "pdf" in pages[p]]
    extracted = [p for p in all_pages if "md" in pages[p] or "json" in pages[p]]
    missing = [p for p in pdf_pages if p not in extracted]
    return {"pdf_pages": len(pdf_pages), "extracted_pages": len(extracted), "missing": missing}


def flatten_json(obj: Any) -> str:
    if isinstance(obj, dict):
        if "raw_response" in obj:
            return str(obj["raw_response"])
        if "rows" in obj and isinstance(obj["rows"], list):
            return json.dumps(obj, ensure_ascii=False)
        return "\n".join(flatten_json(v) for v in obj.values())
    if isinstance(obj, list):
        return "\n".join(flatten_json(x) for x in obj)
    return str(obj)


def extract_html_tables(raw_text: str) -> list[str]:
    return re.findall(r"<table[\s\S]*?</table>", raw_text, flags=re.IGNORECASE)


def parse_html_table(html: str) -> tuple[list[str], list[dict[str, Any]]]:
    if BS4_OK:
        soup = BeautifulSoup(html, "html.parser")
        headers = [th.get_text(" ", strip=True) for th in soup.find_all("th")]
        rows = []
        tr_list = soup.find_all("tr")
        for tr in tr_list:
            cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
            if not cells:
                continue
            if headers and cells == headers:
                continue
            if not headers:
                headers = [f"col_{i+1}" for i in range(len(cells))]
            if len(cells) < len(headers):
                cells = cells + [None] * (len(headers) - len(cells))
            row = {headers[i]: cells[i] for i in range(min(len(headers), len(cells)))}
            rows.append(row)
        return headers, rows

    # fallback via pandas
    try:
        dfs = pd.read_html(html)
        if not dfs:
            return [], []
        df = dfs[0]
        return list(df.columns), df.fillna("").to_dict(orient="records")
    except Exception:
        return [], []


def infer_extraction_type(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ["capillary", "pc", "sw"]):
        return "capillary_pressure"
    if any(k in t for k in ["relative permeability", "krw", "kro", "krg"]):
        return "relative_permeability"
    if any(k in t for k in ["porosity", "permeability", "md"]):
        return "porosity_permeability"
    return "general"


def load_chunks_for_doc(doc_name: str) -> list[dict[str, Any]]:
    pages = S.docs.get(doc_name, {})
    chunks: list[dict[str, Any]] = []
    table_counter = 0
    for page_num in sorted(pages.keys()):
        files = pages[page_num]
        raw = ""
        src = ""
        if "json" in files:
            src = files["json"].name
            try:
                obj = json.loads(files["json"].read_text(encoding="utf-8", errors="ignore"))
                raw = flatten_json(obj)
            except Exception:
                raw = files["json"].read_text(encoding="utf-8", errors="ignore")
        elif "md" in files:
            src = files["md"].name
            raw = files["md"].read_text(encoding="utf-8", errors="ignore")
        if not raw.strip():
            continue

        html_tables = extract_html_tables(raw)
        if html_tables:
            for h in html_tables:
                table_counter += 1
                cols, rows = parse_html_table(h)
                text_for_rag = json.dumps(rows, ensure_ascii=False) if rows else h
                sample_id = ""
                if rows:
                    r0 = rows[0]
                    sample_id = str(r0.get("Sample ID") or r0.get("sample_id") or r0.get("Sample") or "")
                chunks.append(
                    {
                        "text": text_for_rag,
                        "meta": {
                            "file_name": src,
                            "report_name": doc_name,
                            "page_number": page_num,
                            "table_id": f"T{page_num:03d}_{table_counter:02d}",
                            "extraction_type": infer_extraction_type(text_for_rag),
                            "title": f"HTML table page {page_num}",
                            "sample_id": sample_id,
                            "raw_html": h,
                            "parsed_columns": cols,
                            "parsed_rows": rows,
                            "units": {},
                        },
                    }
                )
        else:
            chunks.append(
                {
                    "text": raw,
                    "meta": {
                        "file_name": src,
                        "report_name": doc_name,
                        "page_number": page_num,
                        "table_id": f"P{page_num:03d}_FULL",
                        "extraction_type": infer_extraction_type(raw),
                        "title": f"Full page {page_num}",
                        "sample_id": "",
                        "raw_html": "",
                        "parsed_columns": [],
                        "parsed_rows": [],
                        "units": {},
                    },
                }
            )
    return chunks


def namespace_for_doc(doc_name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", doc_name)


def save_index(namespace: str, vec, mat, texts, metas):
    joblib.dump(vec, INDEX_DIR / f"{namespace}_vec.joblib")
    joblib.dump(mat, INDEX_DIR / f"{namespace}_mat.joblib")
    joblib.dump(texts, INDEX_DIR / f"{namespace}_texts.joblib")
    joblib.dump(metas, INDEX_DIR / f"{namespace}_metas.joblib")


def load_index(namespace: str):
    vec_p = INDEX_DIR / f"{namespace}_vec.joblib"
    mat_p = INDEX_DIR / f"{namespace}_mat.joblib"
    txt_p = INDEX_DIR / f"{namespace}_texts.joblib"
    meta_p = INDEX_DIR / f"{namespace}_metas.joblib"
    if not all(p.exists() for p in [vec_p, mat_p, txt_p, meta_p]):
        return None
    return joblib.load(vec_p), joblib.load(mat_p), joblib.load(txt_p), joblib.load(meta_p)


def build_index_for_current_doc() -> str:
    if not S.current_doc:
        return "No document selected"
    chunks = load_chunks_for_doc(S.current_doc)
    if not chunks:
        return "No extracted chunks (.md/.json) found"
    texts = [c["text"] for c in chunks]
    metas = [c["meta"] for c in chunks]
    vec = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
    mat = vec.fit_transform(texts)
    ns = namespace_for_doc(S.current_doc)
    save_index(ns, vec, mat, texts, metas)

    S.index_namespace = ns
    S.vectorizer, S.matrix = vec, mat
    S.index_texts, S.index_meta = texts, metas
    return f"Index built for {S.current_doc}: {len(texts)} chunks"


def ensure_index_loaded() -> bool:
    if S.vectorizer is not None and S.matrix is not None:
        return True
    if not S.current_doc:
        return False
    ns = namespace_for_doc(S.current_doc)
    obj = load_index(ns)
    if obj is None:
        return False
    S.vectorizer, S.matrix, S.index_texts, S.index_meta = obj
    S.index_namespace = ns
    return True


def search(query: str, filters: dict[str, Any], top_k: int = 8) -> list[dict[str, Any]]:
    if not ensure_index_loaded():
        return []
    qv = S.vectorizer.transform([query])
    sims = linear_kernel(qv, S.matrix).flatten()
    order = sims.argsort()[::-1]
    out = []
    for i in order:
        score = float(sims[i])
        if score <= 0:
            continue
        m = S.index_meta[i]
        ok = True
        for k, v in filters.items():
            if v in (None, ""):
                continue
            if str(m.get(k, "")).lower() != str(v).lower():
                ok = False
                break
        if not ok:
            continue
        out.append({"score": score, "text": S.index_texts[i], "meta": m})
        if len(out) >= top_k:
            break
    return out


def load_prompt_store() -> list[dict[str, str]]:
    if not PROMPT_FILE.exists():
        PROMPT_FILE.write_text(json.dumps(DEFAULT_CHAT_PROMPTS, indent=2), encoding="utf-8")
    try:
        data = json.loads(PROMPT_FILE.read_text(encoding="utf-8"))
        if isinstance(data, list):
            return data
    except Exception:
        pass
    return DEFAULT_CHAT_PROMPTS


def save_prompt_store(prompts: list[dict[str, str]]) -> None:
    PROMPT_FILE.write_text(json.dumps(prompts, indent=2, ensure_ascii=False), encoding="utf-8")


def load_local_llm_runtime(model_name: str):
    import torch
    from transformers import AutoModelForCausalLM, AutoTokenizer

    if not torch.cuda.is_available():
        raise RuntimeError("CUDA GPU required for local LLM")
    tok = AutoTokenizer.from_pretrained(model_name)
    mdl = AutoModelForCausalLM.from_pretrained(model_name, torch_dtype=torch.float16).to("cuda").eval()
    S.llm_tokenizer = tok
    S.llm_model_obj = mdl
    S.llm_model = model_name
    S.llm_loaded = True


def ask_local_llm(system_prompt: str, user_prompt: str) -> str:
    import torch

    if not S.llm_loaded or S.llm_tokenizer is None or S.llm_model_obj is None:
        raise RuntimeError("LLM not loaded")
    msgs = [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]
    inp = S.llm_tokenizer.apply_chat_template(msgs, return_tensors="pt", add_generation_prompt=True).to("cuda", dtype=torch.long)
    with torch.no_grad():
        out = S.llm_model_obj.generate(inp, max_new_tokens=700, temperature=0.2, do_sample=True, top_p=0.9)
    return S.llm_tokenizer.decode(out[0][inp.shape[1] :], skip_special_tokens=True).strip()


def load_vlm_runtime():
    from scal_webapp.backend.services.web_olmocr_runtime import get_vlm

    vlm = get_vlm()
    vlm.load()
    S.vlm = vlm
    S.vlm_loaded = True


def check_uploaded_pdf_missing(pdf_bytes: bytes, doc_name: str) -> dict[str, Any]:
    if doc_name not in S.docs:
        return {"error": "Document not found in extracted dataset"}
    with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(pdf_bytes)
        tmp_path = Path(tmp.name)
    reader = PdfReader(str(tmp_path))
    total = len(reader.pages)
    extracted_pages = {p for p, flags in S.docs[doc_name].items() if "md" in flags or "json" in flags}
    missing = [p for p in range(1, total + 1) if p not in extracted_pages]
    return {
        "total_pdf_pages": total,
        "extracted_pages_found": len(extracted_pages),
        "missing_pages": missing,
        "already_extracted": len(missing) == 0,
    }


def export_results_excel(hits: list[dict[str, Any]]) -> Path:
    rows = []
    for h in hits:
        m = h["meta"]
        rows.append(
            {
                "file_name": m.get("file_name"),
                "report_name": m.get("report_name"),
                "page_number": m.get("page_number"),
                "table_id": m.get("table_id"),
                "extraction_type": m.get("extraction_type"),
                "title": m.get("title"),
                "sample_id": m.get("sample_id"),
                "score": h.get("score"),
                "text": h.get("text"),
            }
        )
    df = pd.DataFrame(rows)
    path = EXPORT_DIR / f"retrieved_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="retrieved", index=False)
    return path


def export_results_word(hits: list[dict[str, Any]]) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("Retrieved SCAL Results", level=1)
    for i, h in enumerate(hits, start=1):
        m = h["meta"]
        doc.add_heading(f"Result {i}", level=2)
        doc.add_paragraph(
            f"File: {m.get('file_name')} | Report: {m.get('report_name')} | "
            f"Page: {m.get('page_number')} | Table: {m.get('table_id')} | Type: {m.get('extraction_type')}"
        )
        if m.get("raw_html") and m.get("parsed_rows"):
            rows = m.get("parsed_rows") or []
            cols = m.get("parsed_columns") or (list(rows[0].keys()) if rows else [])
            if cols:
                t = doc.add_table(rows=1, cols=len(cols))
                hdr = t.rows[0].cells
                for c_idx, c in enumerate(cols):
                    hdr[c_idx].text = str(c)
                for r in rows:
                    tr = t.add_row().cells
                    for c_idx, c in enumerate(cols):
                        tr[c_idx].text = str(r.get(c, ""))
        else:
            doc.add_paragraph(str(h.get("text", ""))[:1800])
    path = EXPORT_DIR / f"retrieved_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(path)
    return path


# ---------------------------- UI ----------------------------

ui.add_head_html("""
<style>
  .oc-shell {height: calc(100vh - 80px);} 
  .oc-chat-scroll {max-height: 62vh; overflow-y: auto;}
</style>
""")

ui.label("SCAL Local Chat Console").classes("text-h5 q-mb-sm")
ui.label("OpenCode-style local UI | JSON-first RAG | optional OLMOCR extraction").classes("text-caption q-mb-md")

with ui.row().classes("w-full items-start"):
    # Left panel
    with ui.card().classes("w-1/4 oc-shell"):
        ui.label("RAG / Data Panel").classes("text-subtitle1")
        data_root_input = ui.input("Extracted data root", value=str(DEFAULT_DATA_ROOT)).classes("w-full")
        doc_select = ui.select(options=[], label="Document").classes("w-full")
        coverage_out = ui.label("Coverage: -")
        missing_out = ui.label("Missing pages: -")

        def refresh_docs():
            S.data_root = Path(data_root_input.value)
            S.docs = scan_dataset(S.data_root)
            names = sorted(S.docs.keys())
            doc_select.options = names
            if names and (S.current_doc not in names):
                S.current_doc = names[0]
                doc_select.value = names[0]
            elif S.current_doc:
                doc_select.value = S.current_doc
            update_coverage()

        def update_coverage():
            if not doc_select.value:
                coverage_out.set_text("Coverage: no document")
                missing_out.set_text("Missing pages: -")
                return
            S.current_doc = doc_select.value
            cov = coverage_info(S.docs.get(S.current_doc, {}))
            coverage_out.set_text(f"PDF pages: {cov['pdf_pages']} | Extracted: {cov['extracted_pages']}")
            missing_out.set_text(f"Missing extracted pages: {cov['missing']}")

        ui.button("Refresh Dataset", on_click=lambda: refresh_docs())
        doc_select.on("update:model-value", lambda _: update_coverage())
        ui.button("Build / Rebuild Index", on_click=lambda: log(build_index_for_current_doc()))

        ui.separator()
        ui.label("Model Controls").classes("text-subtitle2")
        llm_model_input = ui.input("LLM model", value=S.llm_model).classes("w-full")

        def on_load_llm():
            try:
                load_local_llm_runtime(llm_model_input.value)
                model_status.set_text(f"VLM: {'loaded' if S.vlm_loaded else 'not loaded'} | LLM: loaded ({S.llm_model})")
                log("LLM loaded")
            except Exception as e:
                log(f"LLM load failed: {e}")

        def on_load_vlm():
            try:
                load_vlm_runtime()
                model_status.set_text(f"VLM: loaded | LLM: {'loaded' if S.llm_loaded else 'not loaded'}")
                log("VLM loaded")
            except Exception as e:
                log(f"VLM load failed: {e}")

        ui.row().classes("w-full").props('align="between"')
        ui.button("Load VLM", on_click=lambda: on_load_vlm())
        ui.button("Load LLM", on_click=lambda: on_load_llm())

        ui.separator()
        ui.label("Prompt Templates").classes("text-subtitle2")
        prompt_store = load_prompt_store()
        prompt_select = ui.select(options=[p["name"] for p in prompt_store], value=prompt_store[0]["name"]).classes("w-full")
        prompt_area = ui.textarea("Query prompt", value=prompt_store[0]["text"]).classes("w-full")
        prompt_name_input = ui.input("Save prompt as", value=prompt_store[0]["name"]).classes("w-full")

        def on_prompt_select():
            prompts = load_prompt_store()
            for p in prompts:
                if p["name"] == prompt_select.value:
                    prompt_area.value = p["text"]
                    break

        def on_save_prompt():
            prompts = load_prompt_store()
            found = False
            for p in prompts:
                if p["name"] == prompt_name_input.value:
                    p["text"] = prompt_area.value
                    found = True
                    break
            if not found:
                prompts.append({"name": prompt_name_input.value, "text": prompt_area.value})
            save_prompt_store(prompts)
            prompt_select.options = [p["name"] for p in prompts]
            prompt_select.value = prompt_name_input.value
            log(f"Saved prompt: {prompt_name_input.value}")

        prompt_select.on("update:model-value", lambda _: on_prompt_select())
        ui.button("Save Prompt", on_click=lambda: on_save_prompt())

    # Center chat panel
    with ui.card().classes("w-2/4 oc-shell"):
        ui.label("Chat").classes("text-subtitle1")
        model_status = ui.label("VLM: not loaded | LLM: not loaded").classes("text-caption")
        chat_container = ui.column().classes("w-full oc-chat-scroll")

        with ui.row().classes("w-full"):
            file_filter = ui.input("file_name", value="").classes("w-1/5")
            report_filter = ui.input("report_name", value="").classes("w-1/5")
            page_filter = ui.input("page_number", value="").classes("w-1/6")
            table_filter = ui.input("table_id", value="").classes("w-1/6")
            sample_filter = ui.input("sample_id", value="").classes("w-1/6")

        user_input = ui.textarea("Ask question", placeholder="Ask about extracted SCAL data...").classes("w-full")

        def render_chat():
            chat_container.clear()
            history = S.chat_history.get(S.current_session, [])
            with chat_container:
                for q, a in history:
                    with ui.card().classes("w-full bg-blue-1"):
                        ui.label("You").classes("text-caption")
                        ui.markdown(q)
                    with ui.card().classes("w-full bg-grey-2"):
                        ui.label("Assistant").classes("text-caption")
                        ui.markdown(a)

        def handle_chat():
            q = (user_input.value or "").strip()
            if not q:
                return
            if not ensure_index_loaded():
                answer = "No index loaded. Please build index first."
                S.chat_history.setdefault(S.current_session, []).append((q, answer))
                render_chat()
                return

            filters = {
                "file_name": file_filter.value,
                "report_name": report_filter.value,
                "page_number": page_filter.value,
                "table_id": table_filter.value,
                "sample_id": sample_filter.value,
            }
            hits = search(q, filters, top_k=8)
            S.last_hits = hits
            S.last_render_tables = [h for h in hits if h["meta"].get("parsed_rows")]
            refresh_evidence_panel()

            if not hits:
                answer = "No relevant chunks found with current filters."
            else:
                context = []
                for i, h in enumerate(hits, start=1):
                    txt = h["text"]
                    if len(txt) > 700:
                        txt = txt[:700] + "..."
                    m = h["meta"]
                    context.append(
                        f"[{i}] file={m.get('file_name')} page={m.get('page_number')} table={m.get('table_id')}\n{txt}"
                    )
                ctx = "\n\n".join(context)
                system = "You are a SCAL assistant. Use ONLY retrieved evidence. Include citations [1],[2]."
                user_prompt = f"Task prompt:\n{prompt_area.value}\n\nQuestion:\n{q}\n\nRetrieved evidence:\n{ctx}"
                try:
                    answer = ask_local_llm(system, user_prompt)
                except Exception as e:
                    answer = f"Local LLM failed: {e}\n\nFallback evidence:\n{ctx[:1800]}"

            S.chat_history.setdefault(S.current_session, []).append((q, answer))
            user_input.value = ""
            render_chat()
            log("Chat query processed")

        with ui.row().classes("w-full"):
            ui.button("Send", on_click=lambda: handle_chat())
            ui.button("Clear History", on_click=lambda: clear_history())

        def clear_history():
            S.chat_history[S.current_session] = []
            render_chat()
            log("Chat history cleared")

    # Right evidence/export/debug panel
    with ui.card().classes("w-1/4 oc-shell"):
        ui.label("Results / Source Panel").classes("text-subtitle1")
        evidence_json = ui.textarea("Retrieved chunks JSON", value="").props("readonly").classes("w-full")

        table_render_holder = ui.column().classes("w-full")

        def refresh_evidence_panel():
            evidence_json.value = json.dumps(S.last_hits, indent=2, ensure_ascii=False)[:20000]
            table_render_holder.clear()
            with table_render_holder:
                if not S.last_render_tables:
                    ui.label("No parsed HTML table in current retrieval result.").classes("text-caption")
                else:
                    ui.label("Parsed HTML tables").classes("text-subtitle2")
                    for i, h in enumerate(S.last_render_tables[:3], start=1):
                        m = h["meta"]
                        cols = m.get("parsed_columns") or []
                        rows = m.get("parsed_rows") or []
                        ui.label(f"[{i}] {m.get('file_name')} | page {m.get('page_number')} | {m.get('table_id')}").classes("text-caption")
                        if cols and rows:
                            ui.table(
                                columns=[{"name": c, "label": c, "field": c} for c in cols],
                                rows=rows[:20],
                                row_key=cols[0],
                            ).classes("w-full")

        def export_excel_action():
            if not S.last_hits:
                ui.notify("No retrieved results to export", color="warning")
                return
            p = export_results_excel(S.last_hits)
            ui.download(str(p))
            log(f"Excel exported: {p.name}")

        def export_word_action():
            if not S.last_hits:
                ui.notify("No retrieved results to export", color="warning")
                return
            p = export_results_word(S.last_hits)
            ui.download(str(p))
            log(f"Word exported: {p.name}")

        with ui.row().classes("w-full"):
            ui.button("Export Excel", on_click=lambda: export_excel_action())
            ui.button("Export Word", on_click=lambda: export_word_action())

        ui.separator()
        ui.label("Optional PDF Extraction (OLMOCR VLM)").classes("text-subtitle2")
        pdf_upload = ui.upload(label="Upload PDF to extract missing pages", auto_upload=True)
        extract_status = ui.label("No PDF uploaded")
        extraction_prompt_area = ui.textarea("Extraction prompt", value="").classes("w-full")

        def init_default_prompt():
            try:
                from scal_webapp.backend.services.web_olmocr_runtime import default_olmocr_prompt

                extraction_prompt_area.value = default_olmocr_prompt()
            except Exception:
                extraction_prompt_area.value = (
                    "Attached is one page of a document that you must process. Return markdown with front matter..."
                )

        init_default_prompt()

        uploaded_pdf_bytes = {"data": None, "name": None}

        def on_upload(e):
            uploaded_pdf_bytes["data"] = e.content.read()
            uploaded_pdf_bytes["name"] = e.name
            extract_status.set_text(f"Uploaded: {e.name}")
            if S.current_doc:
                try:
                    c = check_uploaded_pdf_missing(uploaded_pdf_bytes["data"], S.current_doc)
                    extract_status.set_text(
                        f"Uploaded: {e.name} | Already extracted: {c.get('already_extracted')} | Missing pages: {c.get('missing_pages')}"
                    )
                except Exception as ex:
                    extract_status.set_text(f"Coverage check error: {ex}")

        pdf_upload.on_upload(on_upload)

        def run_optional_extraction():
            if not uploaded_pdf_bytes["data"]:
                ui.notify("Upload PDF first", color="warning")
                return
            if not S.vlm_loaded or S.vlm is None:
                ui.notify("Load VLM first", color="warning")
                return

            with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(uploaded_pdf_bytes["data"])
                pdf_path = Path(tmp.name)

            reader = PdfReader(str(pdf_path))
            total_pages = len(reader.pages)

            # Decide missing pages if same doc selected, else all pages
            pages_to_extract = list(range(1, total_pages + 1))
            if S.current_doc and S.current_doc in S.docs:
                existing = {p for p, flags in S.docs[S.current_doc].items() if "md" in flags or "json" in flags}
                missing = [p for p in pages_to_extract if p not in existing]
                pages_to_extract = missing or []

            if not pages_to_extract:
                ui.notify("This file appears already fully extracted", color="positive")
                return

            stem = Path(uploaded_pdf_bytes["name"] or "uploaded").stem
            prompt = extraction_prompt_area.value or ""

            extracted_count = 0
            for p in pages_to_extract:
                try:
                    res = S.vlm.extract_page(str(pdf_path), p, prompt=prompt)
                    # save JSON + MD into data root for immediate reindex
                    out_json = S.data_root / f"{stem}_page{p}.json"
                    out_md = S.data_root / f"{stem}_page{p}.md"
                    out_json.write_text(json.dumps(res, indent=2, ensure_ascii=False), encoding="utf-8")
                    out_md.write_text(str(res.get("raw_response", "")), encoding="utf-8")
                    extracted_count += 1
                except Exception as ex:
                    log(f"Extraction failed page {p}: {ex}")

            refresh_docs()
            if stem in S.docs:
                S.current_doc = stem
                doc_select.value = stem
            msg = f"Optional extraction complete: {extracted_count} page(s) saved"
            extract_status.set_text(msg)
            log(msg)

        ui.button("Extract Missing Pages", on_click=lambda: run_optional_extraction())

        ui.separator()
        ui.label("Logs").classes("text-subtitle2")
        log_box = ui.textarea("", value="").props("readonly").classes("w-full")

        def log(message: str):
            line = f"[{ts()}] {message}"
            if log_box.value:
                log_box.value += "\n" + line
            else:
                log_box.value = line

        ui.button("Clear Logs", on_click=lambda: clear_logs())

        def clear_logs():
            log_box.value = ""

        # initialize data
        refresh_docs()
        render_chat()
        refresh_evidence_panel()
        log("Ready")


ui.run(host="0.0.0.0", port=8088, title="SCAL Local Chat Console")
