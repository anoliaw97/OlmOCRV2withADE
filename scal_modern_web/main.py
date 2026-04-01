from __future__ import annotations

import asyncio
import json
import os
import re
import sys
import threading
import time
import traceback
import uuid
import urllib.error
import urllib.request
from collections import deque
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

os.environ.setdefault("HF_HUB_DISABLE_XET", "1")
os.environ.setdefault("HF_HUB_ENABLE_HF_TRANSFER", "1")
os.environ.setdefault("HF_HUB_DOWNLOAD_TIMEOUT", "120")

# Executor for blocking tkinter dialogs (1 thread — dialogs must be serial)
_DIALOG_EXECUTOR = ThreadPoolExecutor(max_workers=1, thread_name_prefix="tkdialog")

# Executor for GPU inference (1 thread — serialises LLM calls, never blocks event loop)
# Extraction runs in its own plain daemon thread (see extraction_job), not here.
_INFERENCE_EXECUTOR = ThreadPoolExecutor(max_workers=1, thread_name_prefix="llm_infer")

# Executor for CPU-bound RAG search (TF-IDF matrix multiply) so chat stays async
_SEARCH_EXECUTOR = ThreadPoolExecutor(max_workers=4, thread_name_prefix="rag_search")

import joblib
import pandas as pd
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel

try:
    from bs4 import BeautifulSoup

    BS4_OK = True
except Exception:
    BS4_OK = False


ROOT = Path(__file__).resolve().parents[1]
DATA_ROOT = Path(
    r"C:\Users\Mining\Downloads\Fine Tunining Datasets-20260318T052420Z-1-001\Fine Tunining Datasets\train"
)
INDEX_DIR = ROOT / "scal_modern_index"
INDEX_DIR.mkdir(parents=True, exist_ok=True)
EXPORT_DIR = ROOT / "scal_modern_exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)
SESSION_FILE = ROOT / "scal_modern_sessions.json"
SETTINGS_FILE = ROOT / "scal_modern_settings.json"
RESULTS_ROOT = ROOT / "results"
INFERENCE_API_URL = os.environ.get("SCAL_INFERENCE_API_URL", "http://127.0.0.1:8010").rstrip("/")

# Large model cache override (to avoid filling C: drive).
# Kimi-K2 is very large; route its Hugging Face cache to D:.
MODEL_CACHE_DIRS: dict[str, Path] = {
    "moonshotai/Kimi-K2.5": Path(r"D:\hf_cache\moonshotai\Kimi-K2.5"),
}

RETRIEVAL_CONFIGS = [
    "tfidf_unigram",
    "tfidf_unigram_bigram",
    "tfidf_with_metadata",
    "tfidf_normalized_table_text",
    "tfidf_text_plus_metadata",
]

PROMPT_TYPES = [
    "direct_answer",
    "structured_answer",
    "compiled_answer",
    "export_ready_answer",
]

LLM_MODEL_OPTIONS = [
    {
        "name": "Qwen/Qwen3-30B-A3B-Instruct-2507",
        "label": "Qwen3-30B-A3B-Instruct (Best overall RAG)",
        "recommended": True,
        "notes": "MoE long-context model (up to ~262K) with strong synthesis for document RAG",
    },
    {
        "name": "deepseek-ai/DeepSeek-R1-Distill-Qwen-32B",
        "label": "DeepSeek-R1-Distill-Qwen-32B (Best reasoning)",
        "recommended": False,
        "notes": "Strong analytical/multi-step reasoning for complex query answering",
    },
    {
        "name": "zai-org/GLM-4-32B-0414",
        "label": "GLM-4-32B-0414 (Open-source alternative)",
        "recommended": False,
        "notes": "Open-weight GLM model with strong coding/tool capabilities",
    },
    {
        "name": "moonshotai/Kimi-K2.5",
        "label": "Kimi-K2.5 (Very high VRAM)",
        "recommended": False,
        "notes": "Open-weight MoE multimodal model; very heavy for single-GPU local inference",
    },
    {
        "name": "Qwen/Qwen2.5-14B-Instruct",
        "label": "Qwen2.5-14B-Instruct (Stable fallback)",
        "recommended": False,
        "notes": "Stable default for lower VRAM pressure and fast iteration",
    },
    {
        "name": "Qwen/Qwen2.5-7B-Instruct",
        "label": "Qwen2.5-7B-Instruct (Faster fallback)",
        "recommended": False,
        "notes": "Lower VRAM and faster responses",
    },
    {
        "name": "meta-llama/Llama-3.1-8B-Instruct",
        "label": "Llama-3.1-8B-Instruct (Alternative)",
        "recommended": False,
        "notes": "Good general instruction model",
    },
]

USE_CASE_PROMPT_SUGGESTIONS = [
    {
        "id": "scal_summary",
        "label": "Summarize SCAL report highlights",
        "question": "Summarize the key SCAL findings for this report with supporting evidence.",
        "prompt_template": "Return concise technical summary grouped by: porosity/permeability, capillary pressure, relative permeability, key sample IDs, and anomalies. Cite sources [1],[2].",
    },
    {
        "id": "extract_por_perm",
        "label": "Extract porosity and permeability table",
        "question": "Extract porosity and permeability values by sample as a table.",
        "prompt_template": "From evidence, return structured table columns: Sample ID, Depth, Porosity, Air Permeability, Grain Density, Notes. Keep units exactly.",
    },
    {
        "id": "extract_cap_pressure",
        "label": "Extract capillary pressure points",
        "question": "Extract capillary pressure vs saturation data points by sample.",
        "prompt_template": "Return JSON array grouped by core/sample with fields: capillary_pressure_psi, liquid_saturation_pct_pv, source_table.",
    },
    {
        "id": "qc_missing_values",
        "label": "QC missing values",
        "question": "Check for missing values, malformed rows, and inconsistent units.",
        "prompt_template": "Perform quality checks on extracted rows and return issue list with severity, row identifier, and suggested correction.",
    },
]


def now() -> str:
    return datetime.now().strftime("%H:%M:%S")


class Runtime:
    def __init__(self):
        self.lock = threading.Lock()
        self.docs: dict[str, dict[int, dict[str, Path]]] = {}
        self.current_doc: str | None = None

        self.vectorizer: TfidfVectorizer | None = None
        self.matrix = None
        self.index_texts: list[str] = []
        self.index_meta: list[dict[str, Any]] = []
        self.index_namespace: str | None = None

        self.progress = {
            "index": {"running": False, "percent": 0, "stage": "idle", "detail": ""},
            "extract": {"running": False, "percent": 0, "stage": "idle", "detail": ""},
            "model": {"running": False, "percent": 0, "stage": "idle", "detail": ""},
        }

        self.logs = {
            "status": deque(maxlen=500),
            "debug": deque(maxlen=500),
            "error": deque(maxlen=500),
        }

        self.llm_loaded = False
        self.vlm_loaded = False
        self.llm_model_name = "Qwen/Qwen2.5-14B-Instruct"
        self.llm_target_model = ""
        self.llm_last_error = ""
        self.last_llm_metrics = {"generation_ms": 0.0, "answer_tokens": 0, "tokens_per_sec": 0.0}
        self._llm_tok = None
        self._llm_model = None
        self._llm_lock = threading.Lock()

        self.vlm = None
        self.extract_stop = threading.Event()  # set to request stop
        self.advanced_mode = False
        self.experiment_stop = threading.Event()
        self.experiment = {
            "running": False,
            "percent": 0,
            "stage": "idle",
            "detail": "",
            "run_id": "",
            "mode": "",
            "output_dir": "",
            "started_at": "",
            "finished_at": "",
            "error": "",
        }
        self.benchmark_stop = threading.Event()
        self.benchmark = {
            "running": False,
            "percent": 0,
            "stage": "idle",
            "detail": "",
            "run_id": "",
            "output_dir": "",
            "started_at": "",
            "finished_at": "",
            "error": "",
            "results": [],
        }


R = Runtime()


def log(kind: str, message: str):
    if kind not in R.logs:
        kind = "debug"
    R.logs[kind].append({"time": now(), "kind": kind, "message": message})


def _load_sessions() -> dict[str, Any]:
    if not SESSION_FILE.exists():
        return {"sessions": []}
    try:
        data = json.loads(SESSION_FILE.read_text(encoding="utf-8"))
        if isinstance(data, dict) and isinstance(data.get("sessions"), list):
            return data
    except Exception:
        pass
    return {"sessions": []}


def _save_sessions(data: dict[str, Any]):
    SESSION_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def _load_settings() -> dict[str, Any]:
    if not SETTINGS_FILE.exists():
        return {"advanced_mode": False}
    try:
        data = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            return {"advanced_mode": bool(data.get("advanced_mode", False))}
    except Exception:
        pass
    return {"advanced_mode": False}


def _save_settings(data: dict[str, Any]):
    SETTINGS_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


R.advanced_mode = bool(_load_settings().get("advanced_mode", False))


def list_sessions() -> list[dict[str, Any]]:
    data = _load_sessions()
    items = []
    for s in data.get("sessions", []):
        items.append(
            {
                "id": s.get("id", ""),
                "title": s.get("title", "Untitled Session"),
                "created_at": s.get("created_at", ""),
                "updated_at": s.get("updated_at", ""),
                "message_count": len(s.get("messages", [])),
            }
        )
    items.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
    return items


def get_session(session_id: str) -> dict[str, Any] | None:
    data = _load_sessions()
    for s in data.get("sessions", []):
        if s.get("id") == session_id:
            return s
    return None


def create_session(title: str = "") -> dict[str, Any]:
    data = _load_sessions()
    sid = str(uuid.uuid4())
    ts = datetime.now().isoformat(timespec="seconds")
    obj = {
        "id": sid,
        "title": title or f"Session {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "created_at": ts,
        "updated_at": ts,
        "messages": [],
    }
    data["sessions"].append(obj)
    _save_sessions(data)
    return obj


def append_session_messages(session_id: str, new_messages: list[dict[str, Any]]):
    data = _load_sessions()
    found = None
    for s in data.get("sessions", []):
        if s.get("id") == session_id:
            found = s
            break
    if found is None:
        found = create_session("Chat Session")
        data = _load_sessions()
        for s in data.get("sessions", []):
            if s.get("id") == found.get("id"):
                found = s
                break
    found.setdefault("messages", []).extend(new_messages)
    found["updated_at"] = datetime.now().isoformat(timespec="seconds")
    if len(found["messages"]) > 300:
        found["messages"] = found["messages"][-300:]
    _save_sessions(data)


def parse_name(file_name: str) -> tuple[str | None, int | None, str]:
    m = re.match(r"^(.*)_page(\d+)\.(pdf|md|json)$", file_name, flags=re.IGNORECASE)
    if not m:
        return None, None, Path(file_name).suffix.lower().lstrip(".")
    return m.group(1), int(m.group(2)), m.group(3).lower()


def scan_docs(root: Path) -> dict[str, dict[int, dict[str, Path]]]:
    """Scan folder for extracted page files AND plain PDFs.

    Recognised patterns (case-insensitive):
      - <stem>_page<N>.pdf / .md / .json  → grouped under <stem>
      - <stem>.pdf (plain, no page number) → added as page 1 under <stem>
    """
    docs: dict[str, dict[int, dict[str, Path]]] = {}
    if not root.exists():
        return docs
    for p in root.iterdir():
        if not p.is_file():
            continue
        stem, page, ext = parse_name(p.name)
        if stem is not None and page is not None:
            # Normal _pageN file
            docs.setdefault(stem, {}).setdefault(page, {})[ext] = p
        elif p.suffix.lower() == ".pdf":
            # Plain PDF — treat as its own single-page document
            doc_stem = p.stem
            docs.setdefault(doc_stem, {}).setdefault(1, {})["pdf"] = p
    return docs


def coverage_for_doc(doc_name: str) -> dict[str, Any]:
    pages = R.docs.get(doc_name, {})
    if not pages:
        return {"pdf_pages": 0, "extracted_pages": 0, "missing_pages": []}
    all_pages = sorted(pages.keys())
    pdf_pages = [p for p in all_pages if "pdf" in pages[p]]
    ext_pages = [p for p in all_pages if "md" in pages[p] or "json" in pages[p]]
    missing = [p for p in pdf_pages if p not in ext_pages]
    return {"pdf_pages": len(pdf_pages), "extracted_pages": len(ext_pages), "missing_pages": missing}


def flatten_json(obj: Any) -> str:
    if isinstance(obj, dict):
        if "raw_response" in obj:
            return str(obj["raw_response"])
        if "rows" in obj and isinstance(obj["rows"], list):
            return json.dumps(obj, ensure_ascii=False)
        return "\n".join(flatten_json(v) for v in obj.values())
    if isinstance(obj, list):
        return "\n".join(flatten_json(x) for x in obj)
    return str(obj)


def extract_html_tables(text: str) -> list[str]:
    return re.findall(r"<table[\s\S]*?</table>", text, flags=re.IGNORECASE)


def parse_html_table(html: str) -> tuple[list[str], list[dict[str, Any]]]:
    if BS4_OK:
        soup = BeautifulSoup(html, "html.parser")
        headers = [th.get_text(" ", strip=True) for th in soup.find_all("th")]
        rows = []
        for tr in soup.find_all("tr"):
            cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
            if not cells:
                continue
            if headers and cells == headers:
                continue
            if not headers:
                headers = [f"col_{i+1}" for i in range(len(cells))]
            if len(cells) < len(headers):
                cells += [None] * (len(headers) - len(cells))
            rows.append({headers[i]: cells[i] for i in range(min(len(headers), len(cells)))})
        return headers, rows
    return [], []


def infer_type(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ["capillary", "pc", "sw"]):
        return "capillary_pressure"
    if any(k in t for k in ["relative permeability", "krw", "kro", "krg"]):
        return "relative_permeability"
    if any(k in t for k in ["porosity", "permeability", "md"]):
        return "porosity_permeability"
    return "general"


def is_casual_chat(query: str) -> bool:
    q = (query or "").strip().lower()
    if not q:
        return False
    simple = {
        "hi",
        "hello",
        "hey",
        "yo",
        "good morning",
        "good afternoon",
        "good evening",
        "thanks",
        "thank you",
        "ok",
        "okay",
        "nice",
        "cool",
        "how are you",
        "who are you",
        "what can you do",
    }
    if q in simple:
        return True
    # Short small-talk style messages should not trigger RAG retrieval
    tokens = re.findall(r"[a-zA-Z0-9']+", q)
    if len(tokens) <= 3 and any(w in q for w in ["hi", "hello", "hey", "thanks", "yo"]):
        return True
    return False


def chunks_for_doc(doc_name: str) -> list[dict[str, Any]]:
    pages = R.docs.get(doc_name, {})
    chunks = []
    tcount = 0
    for pg in sorted(pages.keys()):
        files = pages[pg]
        raw, source = "", ""
        if "json" in files:
            source = files["json"].name
            try:
                raw = flatten_json(json.loads(files["json"].read_text(encoding="utf-8", errors="ignore")))
            except Exception:
                raw = files["json"].read_text(encoding="utf-8", errors="ignore")
        elif "md" in files:
            source = files["md"].name
            raw = files["md"].read_text(encoding="utf-8", errors="ignore")
        if not raw.strip():
            continue
        tables = extract_html_tables(raw)
        if tables:
            for h in tables:
                tcount += 1
                cols, rows = parse_html_table(h)
                txt = json.dumps(rows, ensure_ascii=False) if rows else h
                sample = ""
                if rows:
                    sample = str(rows[0].get("Sample ID") or rows[0].get("sample_id") or "")
                chunks.append(
                    {
                        "text": txt,
                        "meta": {
                            "file_name": source,
                            "report_name": doc_name,
                            "page_number": pg,
                            "table_id": f"T{pg:03d}_{tcount:02d}",
                            "extraction_type": infer_type(txt),
                            "title": f"HTML table page {pg}",
                            "sample_id": sample,
                            "raw_html": h,
                            "parsed_columns": cols,
                            "parsed_rows": rows,
                            "units": {},
                        },
                    }
                )
        else:
            chunks.append(
                {
                    "text": raw,
                    "meta": {
                        "file_name": source,
                        "report_name": doc_name,
                        "page_number": pg,
                        "table_id": f"P{pg:03d}_FULL",
                        "extraction_type": infer_type(raw),
                        "title": f"Full page {pg}",
                        "sample_id": "",
                        "raw_html": "",
                        "parsed_columns": [],
                        "parsed_rows": [],
                        "units": {},
                    },
                }
            )
    return chunks


def ns(doc_name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", doc_name)


def save_index(namespace: str, vec, mat, texts, metas):
    joblib.dump(vec, INDEX_DIR / f"{namespace}_vec.joblib")
    joblib.dump(mat, INDEX_DIR / f"{namespace}_mat.joblib")
    joblib.dump(texts, INDEX_DIR / f"{namespace}_texts.joblib")
    joblib.dump(metas, INDEX_DIR / f"{namespace}_metas.joblib")


def load_index(namespace: str):
    paths = [
        INDEX_DIR / f"{namespace}_vec.joblib",
        INDEX_DIR / f"{namespace}_mat.joblib",
        INDEX_DIR / f"{namespace}_texts.joblib",
        INDEX_DIR / f"{namespace}_metas.joblib",
    ]
    if not all(p.exists() for p in paths):
        return None
    return joblib.load(paths[0]), joblib.load(paths[1]), joblib.load(paths[2]), joblib.load(paths[3])



def set_progress(kind: str, percent: int, stage: str, detail: str = ""):
    if kind not in R.progress:
        return
    R.progress[kind].update({"percent": max(0, min(100, int(percent))), "stage": stage, "detail": detail})


def build_index_job(doc_name: str):
    try:
        R.progress["index"]["running"] = True
        set_progress("index", 5, "scanning", f"Scanning extracted files for {doc_name}")
        log("status", f"Starting index build for {doc_name}")

        chunks = chunks_for_doc(doc_name)
        set_progress("index", 35, "parsing", f"Prepared {len(chunks)} chunks")
        if not chunks:
            log("error", "No extracted chunks found")
            set_progress("index", 100, "failed", "No chunks")
            return

        texts = [c["text"] for c in chunks]
        metas = [c["meta"] for c in chunks]
        vec = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
        mat = vec.fit_transform(texts)
        set_progress("index", 70, "vectorizing", f"Vectorized {len(texts)} chunks")

        namespace = ns(doc_name)
        save_index(namespace, vec, mat, texts, metas)

        R.vectorizer, R.matrix = vec, mat
        R.index_texts, R.index_meta = texts, metas
        R.index_namespace = namespace
        set_progress("index", 100, "completed", f"Index ready ({len(texts)} chunks)")
        log("status", f"Index build completed for {doc_name}")
    except Exception as e:
        log("error", f"Index build failed: {e}")
        log("debug", traceback.format_exc())
        set_progress("index", 100, "failed", str(e))
    finally:
        R.progress["index"]["running"] = False


def chunks_for_docs(doc_names: list[str]) -> list[dict[str, Any]]:
    out = []
    for name in doc_names:
        out.extend(chunks_for_doc(name))
    return out


def build_global_index_job(doc_names: list[str]):
    if not doc_names:
        return 0
    chunks = chunks_for_docs(doc_names)
    if not chunks:
        return 0
    texts = [c["text"] for c in chunks]
    metas = [c["meta"] for c in chunks]
    vec = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
    mat = vec.fit_transform(texts)
    save_index(ns("__ALL__"), vec, mat, texts, metas)
    log("status", f"Global all-doc index updated ({len(texts)} chunks)")
    return len(texts)


def ensure_index_loaded(doc_name: str) -> bool:
    if R.vectorizer is not None and R.matrix is not None and R.current_doc == doc_name:
        return True
    obj = load_index(ns(doc_name))
    if obj is None:
        chunks = chunks_for_doc(doc_name)
        if not chunks:
            return False
        texts = [c["text"] for c in chunks]
        metas = [c["meta"] for c in chunks]
        vec = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
        mat = vec.fit_transform(texts)
        save_index(ns(doc_name), vec, mat, texts, metas)
        obj = (vec, mat, texts, metas)
        log("status", f"Auto-built index for {doc_name} ({len(texts)} chunks)")
    R.vectorizer, R.matrix, R.index_texts, R.index_meta = obj
    R.current_doc = doc_name
    return True


def ensure_all_index_loaded() -> bool:
    obj = load_index(ns("__ALL__"))
    if obj is None:
        names = sorted(R.docs.keys())
        if not names:
            return False
        n = build_global_index_job(names)
        if n <= 0:
            return False
        obj = load_index(ns("__ALL__"))
        if obj is None:
            return False
        log("status", f"Auto-built global index ({n} chunks)")
    R.vectorizer, R.matrix, R.index_texts, R.index_meta = obj
    R.current_doc = "__ALL__"
    return True


def search(query: str, doc_name: str, filters: dict[str, Any], top_k: int = 8) -> list[dict[str, Any]]:
    if doc_name == "__ALL__":
        if not ensure_all_index_loaded():
            # Fallback: aggregate top hits from each available doc index
            merged = []
            for d in sorted(R.docs.keys()):
                if not ensure_index_loaded(d):
                    continue
                qv = R.vectorizer.transform([query])
                sims = linear_kernel(qv, R.matrix).flatten()
                order = sims.argsort()[::-1]
                for i in order[: max(8, top_k)]:
                    score = float(sims[i])
                    if score <= 0:
                        continue
                    m = R.index_meta[i]
                    ok = True
                    for k, v in filters.items():
                        if v in (None, ""):
                            continue
                        if str(m.get(k, "")).lower() != str(v).lower():
                            ok = False
                            break
                    if ok:
                        merged.append({"score": score, "text": R.index_texts[i], "meta": m})
            merged.sort(key=lambda x: x["score"], reverse=True)
            return merged[:top_k]
    elif not ensure_index_loaded(doc_name):
        return []
    qv = R.vectorizer.transform([query])
    sims = linear_kernel(qv, R.matrix).flatten()
    order = sims.argsort()[::-1]
    out = []
    for i in order:
        score = float(sims[i])
        if score <= 0:
            continue
        m = R.index_meta[i]
        ok = True
        for k, v in filters.items():
            if v in (None, ""):
                continue
            if str(m.get(k, "")).lower() != str(v).lower():
                ok = False
                break
        if not ok:
            continue
        out.append({"score": score, "text": R.index_texts[i], "meta": m})
        if len(out) >= top_k:
            break
    return out


def _inference_request(method: str, path: str, payload: dict[str, Any] | None = None, timeout: int = 20) -> dict[str, Any]:
    url = f"{INFERENCE_API_URL}{path}"
    data = None
    headers = {"Accept": "application/json"}
    if payload is not None:
        data = json.dumps(payload).encode("utf-8")
        headers["Content-Type"] = "application/json"
    req = urllib.request.Request(url, data=data, headers=headers, method=method.upper())
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            body = resp.read().decode("utf-8", errors="ignore")
            return json.loads(body) if body else {}
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"Inference API HTTP {e.code}: {detail or e.reason}")
    except Exception as e:
        raise RuntimeError(f"Inference API unavailable at {INFERENCE_API_URL}: {e}")


def _inference_stream_events(path: str, payload: dict[str, Any], timeout: int = 600):
    url = f"{INFERENCE_API_URL}{path}"
    data = json.dumps(payload).encode("utf-8")
    headers = {"Accept": "text/event-stream", "Content-Type": "application/json"}
    req = urllib.request.Request(url, data=data, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            for raw in resp:
                line = raw.decode("utf-8", errors="ignore").strip()
                if not line.startswith("data:"):
                    continue
                data_line = line[5:].strip()
                if not data_line:
                    continue
                if data_line == "[DONE]":
                    break
                try:
                    yield json.loads(data_line)
                except Exception:
                    continue
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"Inference API HTTP {e.code}: {detail or e.reason}")
    except Exception as e:
        raise RuntimeError(f"Inference API stream unavailable at {INFERENCE_API_URL}: {e}")


def sync_llm_state_from_inference() -> dict[str, Any]:
    try:
        st = _inference_request("GET", "/v1/health", timeout=5)
        R.llm_loaded = bool(st.get("loaded", False))
        R.llm_model_name = str(st.get("model_name") or R.llm_model_name)
        R.llm_target_model = str(st.get("target_model") or "")
        R.llm_last_error = str(st.get("last_error") or "")
        p = st.get("progress") or {}
        R.progress["model"] = {
            "running": bool(st.get("busy", False) or st.get("state") in {"loading", "unloading"}),
            "percent": int(p.get("percent", 0) or 0),
            "stage": str(p.get("stage", st.get("state", "idle"))),
            "detail": str(p.get("detail", "")),
        }
        return st
    except Exception as e:
        R.llm_loaded = False
        R.llm_target_model = ""
        R.llm_last_error = str(e)
        R.progress["model"].update({"running": False, "stage": "failed", "detail": str(e)})
        return {}


def load_llm(model_name: str):
    R.llm_target_model = model_name
    R.llm_last_error = ""
    R.progress["model"]["running"] = True
    set_progress("model", 5, "starting", f"Requesting model load: {model_name}")
    resp = _inference_request("POST", "/v1/models/load", {"model_name": model_name}, timeout=15)
    if not resp.get("ok", False):
        raise RuntimeError(str(resp.get("message") or "Model load request rejected"))
    sync_llm_state_from_inference()
    return resp


def unload_llm():
    R.progress["model"]["running"] = True
    set_progress("model", 10, "unloading", "Requesting model unload")
    resp = _inference_request("POST", "/v1/models/unload", {}, timeout=15)
    if not resp.get("ok", False):
        raise RuntimeError(str(resp.get("message") or "Model unload request rejected"))
    sync_llm_state_from_inference()
    return resp


def ask_llm(
    system_prompt: str,
    user_prompt: str,
    max_new_tokens: int = 420,
    temperature: float = 0.2,
    top_p: float = 0.9,
    do_sample: bool = True,
) -> str:
    payload = {
        "system_prompt": system_prompt,
        "user_prompt": user_prompt,
        "max_new_tokens": int(max_new_tokens),
        "temperature": float(temperature),
        "top_p": float(top_p),
        "do_sample": bool(do_sample),
    }
    resp = _inference_request("POST", "/v1/chat/completions", payload, timeout=600)
    metrics = resp.get("metrics") or {}
    R.last_llm_metrics = {
        "generation_ms": float(metrics.get("generation_ms", 0) or 0),
        "answer_tokens": int(metrics.get("answer_tokens", 0) or 0),
        "tokens_per_sec": float(metrics.get("tokens_per_sec", 0) or 0),
    }
    return str(resp.get("answer") or "")


def llm_generation_settings(mode: str) -> dict[str, Any]:
    m = (mode or "balanced").strip().lower()
    if m == "fast":
        return {"max_new_tokens": 220, "temperature": 0.0, "top_p": 1.0, "do_sample": False}
    if m == "deep":
        return {"max_new_tokens": 700, "temperature": 0.3, "top_p": 0.95, "do_sample": True}
    return {"max_new_tokens": 420, "temperature": 0.2, "top_p": 0.9, "do_sample": True}


def approx_token_count(text: str) -> int:
    t = text or ""
    if not t:
        return 0
    return len(re.findall(r"\w+|[^\w\s]", t, flags=re.UNICODE))


def load_vlm():
    import sys

    sys.path.insert(0, str(ROOT))
    from scal_webapp.backend.services.web_olmocr_runtime import get_vlm

    v = get_vlm()
    v.load()
    R.vlm = v
    R.vlm_loaded = True


def _slug(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", (name or "").strip())[:80] or "item"


def _set_experiment_progress(percent: int, stage: str, detail: str = ""):
    R.experiment.update(
        {
            "percent": max(0, min(100, int(percent))),
            "stage": stage,
            "detail": detail,
        }
    )


def _load_benchmark(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        raise RuntimeError(f"Benchmark file not found: {path}")
    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    elif path.suffix.lower() in {".json", ".jsonl"}:
        txt = path.read_text(encoding="utf-8", errors="ignore")
        data = json.loads(txt)
        if isinstance(data, dict):
            data = data.get("queries", [])
        df = pd.DataFrame(data)
    else:
        raise RuntimeError("Benchmark must be .csv or .json")

    required = ["query_id", "query_text", "expected_document", "expected_page", "expected_table"]
    for col in required:
        if col not in df.columns:
            raise RuntimeError(f"Benchmark missing required column: {col}")
    return df.fillna("").to_dict(orient="records")


def _collect_units(data_root: Path) -> list[dict[str, Any]]:
    docs = scan_docs(data_root)
    units: list[dict[str, Any]] = []
    for doc_name in sorted(docs.keys()):
        pages = docs.get(doc_name, {})
        for pg in sorted(pages.keys()):
            files = pages[pg]
            raw, source = "", ""
            if "json" in files:
                source = files["json"].name
                try:
                    raw = flatten_json(json.loads(files["json"].read_text(encoding="utf-8", errors="ignore")))
                except Exception:
                    raw = files["json"].read_text(encoding="utf-8", errors="ignore")
            elif "md" in files:
                source = files["md"].name
                raw = files["md"].read_text(encoding="utf-8", errors="ignore")
            if not raw.strip():
                continue
            tables = extract_html_tables(raw)
            if tables:
                for idx, h in enumerate(tables, start=1):
                    cols, rows = parse_html_table(h)
                    txt = json.dumps(rows, ensure_ascii=False) if rows else h
                    table_id = f"T{pg:03d}_{idx:02d}"
                    units.append(
                        {
                            "chunk_id": f"{doc_name}|{pg}|{table_id}",
                            "text": txt,
                            "meta": {
                                "report_name": doc_name,
                                "file_name": source,
                                "page_number": pg,
                                "table_id": table_id,
                                "title": f"HTML table page {pg}",
                                "extraction_type": infer_type(txt),
                            },
                            "chunk_type": "table",
                        }
                    )
            units.append(
                {
                    "chunk_id": f"{doc_name}|{pg}|P{pg:03d}_FULL",
                    "text": raw,
                    "meta": {
                        "report_name": doc_name,
                        "file_name": source,
                        "page_number": pg,
                        "table_id": f"P{pg:03d}_FULL",
                        "title": f"Full page {pg}",
                        "extraction_type": infer_type(raw),
                    },
                    "chunk_type": "page",
                }
            )
    return units


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").lower()).strip()


def _unit_text_for_config(unit: dict[str, Any], config_name: str) -> str:
    text = unit.get("text", "")
    m = unit.get("meta", {})
    meta_txt = " ".join(
        [
            str(m.get("report_name", "")),
            str(m.get("file_name", "")),
            str(m.get("title", "")),
            str(m.get("table_id", "")),
            str(m.get("extraction_type", "")),
        ]
    )
    if config_name == "tfidf_unigram":
        return text
    if config_name == "tfidf_unigram_bigram":
        return text
    if config_name == "tfidf_with_metadata":
        return f"{text}\n{meta_txt}"
    if config_name == "tfidf_normalized_table_text":
        if unit.get("chunk_type") == "table":
            return _normalize_text(text)
        return text
    if config_name == "tfidf_text_plus_metadata":
        return f"{_normalize_text(text)}\n{_normalize_text(meta_txt)}"
    return text


def _build_retrieval_index(units: list[dict[str, Any]], config_name: str):
    texts = [_unit_text_for_config(u, config_name) for u in units]
    ngram = (1, 1) if config_name == "tfidf_unigram" else (1, 2)
    vec = TfidfVectorizer(ngram_range=ngram, min_df=1)
    mat = vec.fit_transform(texts)
    return vec, mat, texts


def _retrieve_topk(vec, mat, texts: list[str], units: list[dict[str, Any]], query: str, config_name: str, top_k: int):
    qtext = query if config_name != "tfidf_text_plus_metadata" else _normalize_text(query)
    qv = vec.transform([qtext])
    sims = linear_kernel(qv, mat).flatten()
    order = sims.argsort()[::-1]
    out = []
    for i in order[: max(20, top_k * 2)]:
        score = float(sims[i])
        if score <= 0:
            continue
        u = units[i]
        out.append(
            {
                "score": score,
                "chunk_id": u.get("chunk_id", ""),
                "report_name": u.get("meta", {}).get("report_name", ""),
                "page_number": u.get("meta", {}).get("page_number", ""),
                "table_id": u.get("meta", {}).get("table_id", ""),
                "text": u.get("text", ""),
            }
        )
        if len(out) >= top_k:
            break
    return out


def _truth_match(hit: dict[str, Any], row: dict[str, Any]) -> tuple[bool, bool, bool, bool]:
    exp_doc = str(row.get("expected_document", "")).strip()
    exp_page = str(row.get("expected_page", "")).strip()
    exp_table = str(row.get("expected_table", "")).strip()
    doc_ok = (not exp_doc) or (str(hit.get("report_name", "")).strip() == exp_doc)
    page_ok = (not exp_page) or (str(hit.get("page_number", "")).strip() == exp_page)
    table_ok = (not exp_table) or (str(hit.get("table_id", "")).strip() == exp_table)
    all_ok = doc_ok and page_ok and table_ok
    return doc_ok, page_ok, table_ok, all_ok


def _evaluate_retrieval(run_dir: Path, units: list[dict[str, Any]], benchmark_rows: list[dict[str, Any]], configs: list[str], top_k: int):
    retrieval_root = run_dir / "retrieval"
    retrieval_root.mkdir(parents=True, exist_ok=True)
    overall = []
    for ci, cfg in enumerate(configs, start=1):
        cdir = retrieval_root / f"config_{ci:02d}"
        cdir.mkdir(parents=True, exist_ok=True)
        vec, mat, _ = _build_retrieval_index(units, cfg)
        rows = []
        fail = 0
        for qidx, row in enumerate(benchmark_rows, start=1):
            if R.experiment_stop.is_set():
                raise RuntimeError("Experiment stopped by user")
            t0 = time.perf_counter()
            hits = _retrieve_topk(vec, mat, [], units, str(row.get("query_text", "")), cfg, max(3, top_k))
            latency = round(time.perf_counter() - t0, 4)
            top1 = hits[0] if hits else {}
            d1, p1, tb1, all1 = _truth_match(top1, row) if hits else (False, False, False, False)
            hit_top3 = False
            for h in hits[:3]:
                *_tmp, all_ok = _truth_match(h, row)
                if all_ok:
                    hit_top3 = True
                    break
            if not hits:
                err = "no_retrieval"
            elif not d1:
                err = "wrong_document"
            elif d1 and not p1:
                err = "wrong_page"
            elif d1 and p1 and not tb1:
                err = "wrong_table"
            elif hit_top3 and not all1:
                err = "partial_match"
            elif all1:
                err = ""
            else:
                err = "noisy_retrieval"
            if err:
                fail += 1
            rows.append(
                {
                    "query_id": row.get("query_id", qidx),
                    "retrieval_configuration_name": cfg,
                    "top_1_result": top1.get("chunk_id", ""),
                    "top_3_results": json.dumps([h.get("chunk_id", "") for h in hits[:3]], ensure_ascii=False),
                    "correct_document_top1": int(d1),
                    "correct_page_top1": int(p1),
                    "correct_table_top1": int(tb1),
                    "hit_in_top3": int(hit_top3),
                    "retrieval_latency": latency,
                    "error_type": err,
                }
            )
            _set_experiment_progress(20 + int((qidx / max(1, len(benchmark_rows))) * 35), "retrieval_eval", f"{cfg}: {qidx}/{len(benchmark_rows)}")

        pq = pd.DataFrame(rows)
        pq.to_csv(cdir / "per_query_results.csv", index=False)
        summary = {
            "retrieval_configuration_name": cfg,
            "top_1_accuracy": round(float((pq["correct_document_top1"] & pq["correct_page_top1"] & pq["correct_table_top1"]).mean()), 4),
            "top_3_recall": round(float(pq["hit_in_top3"].mean()), 4),
            "document_accuracy": round(float(pq["correct_document_top1"].mean()), 4),
            "page_accuracy": round(float(pq["correct_page_top1"].mean()), 4),
            "table_accuracy": round(float(pq["correct_table_top1"].mean()), 4),
            "average_latency": round(float(pq["retrieval_latency"].mean()), 4),
            "total_queries": int(len(pq)),
            "total_failures": int(fail),
        }
        pd.DataFrame([summary]).to_csv(cdir / "summary_metrics.csv", index=False)
        err_df = pq[pq["error_type"].astype(str) != ""].groupby("error_type", as_index=False).size()
        err_df.columns = ["error_type", "count"]
        err_df.to_csv(cdir / "error_breakdown.csv", index=False)
        overall.append(summary)

    ov = pd.DataFrame(overall)
    ov.to_csv(retrieval_root / "overall_retrieval_comparison.csv", index=False)
    return overall


def _best_retrieval_config(run_dir: Path) -> str:
    p = run_dir / "retrieval" / "overall_retrieval_comparison.csv"
    if not p.exists():
        raise RuntimeError("Retrieval comparison file not found")
    df = pd.read_csv(p)
    if df.empty:
        raise RuntimeError("No retrieval summary rows")
    df = df.sort_values(["top_1_accuracy", "top_3_recall", "average_latency"], ascending=[False, False, True])
    return str(df.iloc[0]["retrieval_configuration_name"])


def _prompt_for_type(prompt_type: str, question: str, context: str) -> tuple[str, str]:
    system = "You are a SCAL assistant. Use only provided context and cite [1],[2] when possible."
    if prompt_type == "structured_answer":
        user = f"Question: {question}\n\nReturn JSON with keys: answer, evidence_points, assumptions.\n\nContext:\n{context}"
    elif prompt_type == "compiled_answer":
        user = f"Question: {question}\n\nReturn concise compiled technical answer with bullets and citations.\n\nContext:\n{context}"
    elif prompt_type == "export_ready_answer":
        user = f"Question: {question}\n\nReturn export-ready markdown with section headers and a compact table if useful.\n\nContext:\n{context}"
    else:
        user = f"Question: {question}\n\nAnswer directly and concisely.\n\nContext:\n{context}"
    return system, user


def _score_reasoning(row: dict[str, Any], response: str, prompt_type: str, context_ids: list[str]) -> dict[str, Any]:
    exp_answer = str(row.get("expected_answer", "")).strip().lower()
    keys_raw = str(row.get("expected_keywords", "")).strip()
    keys = [k.strip().lower() for k in keys_raw.split(",") if k.strip()]
    rlow = (response or "").lower()
    if keys:
        hit = sum(1 for k in keys if k in rlow)
        correctness = min(1.0, hit / max(1, len(keys)))
        completeness = correctness
    elif exp_answer:
        correctness = 1.0 if exp_answer in rlow else 0.0
        completeness = correctness
    else:
        correctness = 0.5 if response.strip() else 0.0
        completeness = correctness

    if prompt_type == "structured_answer":
        formatting = 1.0 if response.strip().startswith("{") else 0.0
    elif prompt_type == "export_ready_answer":
        formatting = 1.0 if ("|" in response or "##" in response) else 0.0
    else:
        formatting = 1.0 if response.strip() else 0.0

    refs = [int(x) for x in re.findall(r"\[(\d+)\]", response or "")]
    halluc = int(any(r > max(1, len(context_ids)) for r in refs))
    unnecessary = int(prompt_type == "direct_answer" and len(response or "") > 1200)
    export_ready = int(prompt_type == "export_ready_answer" and formatting >= 1.0 and halluc == 0)
    return {
        "correctness_score": round(correctness, 4),
        "completeness_score": round(completeness, 4),
        "formatting_score": round(formatting, 4),
        "hallucination_flag": halluc,
        "unnecessary_reasoning_flag": unnecessary,
        "export_ready_flag": export_ready,
    }


def _run_reasoning(
    run_dir: Path,
    units: list[dict[str, Any]],
    benchmark_rows: list[dict[str, Any]],
    selected_cfg: str,
    models: list[str],
    prompt_types: list[str],
    top_k: int,
):
    reasoning_root = run_dir / "reasoning"
    reasoning_root.mkdir(parents=True, exist_ok=True)
    vec, mat, _ = _build_retrieval_index(units, selected_cfg)
    overall = []
    total_ops = max(1, len(models) * len(prompt_types) * len(benchmark_rows))
    done = 0

    for model_name in models:
        if R.experiment_stop.is_set():
            raise RuntimeError("Experiment stopped by user")
        model_slug = _slug(model_name)
        mdir = reasoning_root / model_slug
        mdir.mkdir(parents=True, exist_ok=True)
        try:
            load_llm(model_name)
        except Exception as e:
            log("error", f"Reasoning skip model {model_name}: {e}")
            continue

        for prompt_type in prompt_types:
            pdir = mdir / prompt_type
            pdir.mkdir(parents=True, exist_ok=True)
            rows = []
            raw_outputs = []
            for qidx, row in enumerate(benchmark_rows, start=1):
                if R.experiment_stop.is_set():
                    raise RuntimeError("Experiment stopped by user")
                hits = _retrieve_topk(vec, mat, [], units, str(row.get("query_text", "")), selected_cfg, top_k)
                context_ids = [h.get("chunk_id", "") for h in hits]
                ctx = []
                for i, h in enumerate(hits, start=1):
                    txt = str(h.get("text", ""))
                    if len(txt) > 500:
                        txt = txt[:500] + "..."
                    ctx.append(f"[{i}] {h.get('chunk_id')}\n{txt}")
                context = "\n\n".join(ctx)
                system, user = _prompt_for_type(prompt_type, str(row.get("query_text", "")), context)
                t0 = time.perf_counter()
                try:
                    resp = ask_llm(system, user)
                except Exception as e:
                    resp = f"ERROR: {e}"
                latency = round(time.perf_counter() - t0, 4)
                score = _score_reasoning(row, resp, prompt_type, context_ids)
                out_row = {
                    "query_id": row.get("query_id", qidx),
                    "model_name": model_name,
                    "prompt_type": prompt_type,
                    "retrieved_context_ids": json.dumps(context_ids, ensure_ascii=False),
                    "response_text": resp,
                    "response_latency": latency,
                    **score,
                }
                rows.append(out_row)
                raw_outputs.append(
                    {
                        "query_id": row.get("query_id", qidx),
                        "query_text": row.get("query_text", ""),
                        "prompt_type": prompt_type,
                        "retrieved_context_ids": context_ids,
                        "response_text": resp,
                    }
                )
                done += 1
                _set_experiment_progress(60 + int((done / total_ops) * 35), "reasoning_eval", f"{model_slug}/{prompt_type}: {qidx}/{len(benchmark_rows)}")

            pq = pd.DataFrame(rows)
            pq.to_csv(pdir / "per_query_results.csv", index=False)
            summary = {
                "model_name": model_name,
                "prompt_type": prompt_type,
                "average_correctness": round(float(pq["correctness_score"].mean()), 4),
                "average_completeness": round(float(pq["completeness_score"].mean()), 4),
                "average_formatting": round(float(pq["formatting_score"].mean()), 4),
                "hallucination_rate": round(float(pq["hallucination_flag"].mean()), 4),
                "unnecessary_reasoning_rate": round(float(pq["unnecessary_reasoning_flag"].mean()), 4),
                "export_ready_rate": round(float(pq["export_ready_flag"].mean()), 4),
                "average_latency": round(float(pq["response_latency"].mean()), 4),
                "total_queries": int(len(pq)),
            }
            pd.DataFrame([summary]).to_csv(pdir / "summary_metrics.csv", index=False)
            (pdir / "raw_outputs.json").write_text(json.dumps(raw_outputs, indent=2, ensure_ascii=False), encoding="utf-8")
            overall.append(summary)

    pd.DataFrame(overall).to_csv(reasoning_root / "overall_reasoning_comparison.csv", index=False)
    return overall


def _finalize_summary(run_dir: Path, selected_cfg: str):
    final_dir = run_dir / "final_summary"
    final_dir.mkdir(parents=True, exist_ok=True)
    rcmp = pd.read_csv(run_dir / "retrieval" / "overall_retrieval_comparison.csv")
    ycmp = pd.read_csv(run_dir / "reasoning" / "overall_reasoning_comparison.csv")
    best_reasoning = ycmp.sort_values(["average_correctness", "average_completeness", "average_latency"], ascending=[False, False, True]).iloc[0]
    best_retrieval = rcmp[rcmp["retrieval_configuration_name"] == selected_cfg].iloc[0]
    row = {
        "best_retrieval_configuration": selected_cfg,
        "best_llm_model": best_reasoning["model_name"],
        "best_prompt_type": best_reasoning["prompt_type"],
        "retrieval_accuracy": best_retrieval["top_1_accuracy"],
        "reasoning_accuracy": best_reasoning["average_correctness"],
        "retrieval_latency": best_retrieval["average_latency"],
        "reasoning_latency": best_reasoning["average_latency"],
        "overall_recommended_pipeline": f"{selected_cfg} + {best_reasoning['model_name']} + {best_reasoning['prompt_type']}",
    }
    pd.DataFrame([row]).to_csv(final_dir / "best_pipeline_summary.csv", index=False)
    overview = {
        "run_id": run_dir.name,
        "best_retrieval_configuration": selected_cfg,
        "best_llm_model": str(best_reasoning["model_name"]),
        "best_prompt_type": str(best_reasoning["prompt_type"]),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
    }
    (final_dir / "experiment_overview.json").write_text(json.dumps(overview, indent=2, ensure_ascii=False), encoding="utf-8")


def _experiment_worker(payload: dict[str, Any]):
    try:
        R.experiment_stop.clear()
        mode = str(payload.get("mode", "full")).strip().lower()
        run_id = str(payload.get("run_id") or datetime.now().strftime("run_%Y%m%d_%H%M%S"))
        data_root = Path(payload.get("data_root") or DATA_ROOT)
        benchmark_path = Path(payload.get("benchmark_path") or "")
        output_root = Path(payload.get("output_root") or RESULTS_ROOT)
        top_k = max(1, int(payload.get("top_k") or 3))
        run_dir = output_root / "runs" / _slug(run_id)
        run_dir.mkdir(parents=True, exist_ok=True)

        R.experiment.update(
            {
                "running": True,
                "percent": 1,
                "stage": "starting",
                "detail": "Initializing experiment",
                "run_id": run_dir.name,
                "mode": mode,
                "output_dir": str(run_dir),
                "started_at": datetime.now().isoformat(timespec="seconds"),
                "finished_at": "",
                "error": "",
            }
        )

        configs = payload.get("retrieval_configs") or RETRIEVAL_CONFIGS
        configs = [c for c in configs if c in RETRIEVAL_CONFIGS] or RETRIEVAL_CONFIGS
        models = payload.get("model_names") or [m["name"] for m in LLM_MODEL_OPTIONS if "moonshotai/Kimi" not in m["name"]]
        prompt_types = payload.get("prompt_types") or PROMPT_TYPES
        prompt_types = [p for p in prompt_types if p in PROMPT_TYPES] or PROMPT_TYPES

        _set_experiment_progress(5, "loading_data", f"Loading extracted JSON/MD from {data_root}")
        units = _collect_units(data_root)
        if not units:
            raise RuntimeError("No retrieval units found in data root")

        _set_experiment_progress(12, "loading_benchmark", f"Loading benchmark: {benchmark_path}")
        benchmark_rows = _load_benchmark(benchmark_path)
        if not benchmark_rows:
            raise RuntimeError("Benchmark file has no rows")

        retrieval_done = (run_dir / "retrieval" / "overall_retrieval_comparison.csv").exists()
        selected_cfg = str(payload.get("selected_retrieval_config") or "").strip()

        if mode in {"full", "retrieval"}:
            _set_experiment_progress(15, "retrieval_build", "Running retrieval study")
            _evaluate_retrieval(run_dir, units, benchmark_rows, configs, top_k)
            retrieval_done = True

        if mode in {"full", "reasoning"}:
            if not retrieval_done and not selected_cfg:
                raise RuntimeError("Reasoning mode requires retrieval outputs or selected retrieval config")
            if not selected_cfg:
                selected_cfg = _best_retrieval_config(run_dir)
            _set_experiment_progress(60, "reasoning_start", f"Using retrieval config: {selected_cfg}")
            _run_reasoning(run_dir, units, benchmark_rows, selected_cfg, models, prompt_types, top_k)

        if mode == "full":
            _set_experiment_progress(96, "final_summary", "Building final comparison summary")
            _finalize_summary(run_dir, selected_cfg)

        registry = output_root / "run_registry.csv"
        reg_row = {
            "run_id": run_dir.name,
            "mode": mode,
            "started_at": R.experiment.get("started_at", ""),
            "finished_at": datetime.now().isoformat(timespec="seconds"),
            "output_dir": str(run_dir),
        }
        if registry.exists():
            reg_df = pd.read_csv(registry)
            reg_df = pd.concat([reg_df, pd.DataFrame([reg_row])], ignore_index=True)
        else:
            reg_df = pd.DataFrame([reg_row])
        reg_df.to_csv(registry, index=False)

        _set_experiment_progress(100, "completed", f"Experiment completed: {run_dir}")
        R.experiment["finished_at"] = datetime.now().isoformat(timespec="seconds")
    except Exception as e:
        _set_experiment_progress(100, "failed", str(e))
        R.experiment["error"] = str(e)
        R.experiment["finished_at"] = datetime.now().isoformat(timespec="seconds")
        log("error", f"Experiment failed: {e}")
        log("debug", traceback.format_exc())
    finally:
        R.experiment["running"] = False


def _keyword_score(text: str, expected_keywords: str) -> float:
    keys = [k.strip().lower() for k in str(expected_keywords or "").split(",") if k.strip()]
    if not keys:
        return 0.0
    low = (text or "").lower()
    hit = sum(1 for k in keys if k in low)
    return round(hit / max(1, len(keys)), 4)


def _run_model_benchmark_worker(payload: dict[str, Any]):
    try:
        R.benchmark_stop.clear()
        run_id = str(payload.get("run_id") or datetime.now().strftime("bench_%Y%m%d_%H%M%S"))
        output_root = Path(payload.get("output_root") or RESULTS_ROOT)
        run_dir = output_root / "benchmarks" / _slug(run_id)
        run_dir.mkdir(parents=True, exist_ok=True)
        question = str(payload.get("question") or "").strip()
        if not question:
            raise RuntimeError("Question is required")

        models = payload.get("model_names") or [m["name"] for m in LLM_MODEL_OPTIONS if "moonshotai/Kimi" not in m["name"]]
        simple_ctx = str(payload.get("simple_context") or "").strip()
        detailed_ctx = str(payload.get("detailed_context") or "").strip()
        expected = str(payload.get("expected_keywords") or "").strip()

        R.benchmark.update(
            {
                "running": True,
                "percent": 1,
                "stage": "starting",
                "detail": "Preparing model benchmark",
                "run_id": _slug(run_id),
                "output_dir": str(run_dir),
                "started_at": datetime.now().isoformat(timespec="seconds"),
                "finished_at": "",
                "error": "",
                "results": [],
            }
        )

        results = []
        total = max(1, len(models))
        for i, model_name in enumerate(models, start=1):
            if R.benchmark_stop.is_set():
                raise RuntimeError("Benchmark stopped by user")

            R.benchmark.update(
                {
                    "percent": int((i - 1) / total * 100),
                    "stage": "loading_model",
                    "detail": f"Loading {model_name}",
                }
            )
            t_load0 = time.perf_counter()
            load_llm(model_name)
            load_sec = round(time.perf_counter() - t_load0, 4)

            def _ask(ctx: str):
                system = "You are a SCAL assistant. Use only provided context and answer clearly."
                user = f"Context:\n{ctx or '(none)'}\n\nQuestion:\n{question}"
                t0 = time.perf_counter()
                ans = ask_llm(system, user)
                return ans, round(time.perf_counter() - t0, 4)

            R.benchmark.update({"stage": "running_simple", "detail": f"{model_name}: simple context"})
            simple_ans, simple_lat = _ask(simple_ctx)
            simple_score = _keyword_score(simple_ans, expected)

            R.benchmark.update({"stage": "running_detailed", "detail": f"{model_name}: detailed context"})
            detailed_ans, detailed_lat = _ask(detailed_ctx)
            detailed_score = _keyword_score(detailed_ans, expected)

            row = {
                "model_name": model_name,
                "load_seconds": load_sec,
                "simple_latency": simple_lat,
                "detailed_latency": detailed_lat,
                "simple_keyword_score": simple_score,
                "detailed_keyword_score": detailed_score,
                "score_gain_detailed_minus_simple": round(detailed_score - simple_score, 4),
                "simple_preview": simple_ans[:280],
                "detailed_preview": detailed_ans[:280],
            }
            results.append(row)
            R.benchmark["results"] = results[-10:]
            R.benchmark["percent"] = int(i / total * 100)

        df = pd.DataFrame(results)
        df.to_csv(run_dir / "per_model_results.csv", index=False)
        summary = df.sort_values(["detailed_keyword_score", "detailed_latency"], ascending=[False, True])
        summary.to_csv(run_dir / "model_ranking.csv", index=False)
        (run_dir / "benchmark_inputs.json").write_text(
            json.dumps(
                {
                    "question": question,
                    "simple_context": simple_ctx,
                    "detailed_context": detailed_ctx,
                    "expected_keywords": expected,
                    "models": models,
                },
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        R.benchmark.update(
            {
                "running": False,
                "percent": 100,
                "stage": "completed",
                "detail": f"Benchmark done: {run_dir}",
                "finished_at": datetime.now().isoformat(timespec="seconds"),
                "results": results,
            }
        )
    except Exception as e:
        R.benchmark.update(
            {
                "running": False,
                "percent": 100,
                "stage": "failed",
                "detail": str(e),
                "error": str(e),
                "finished_at": datetime.now().isoformat(timespec="seconds"),
            }
        )
        log("error", f"Model benchmark failed: {e}")
        log("debug", traceback.format_exc())



def extraction_job(
    pdf_path: Path,
    stem: str,
    prompt: str,
    doc_name: str | None,
    output_dir: Path | None,
    page_from: int,
    page_to: int,
    is_tmp: bool = False,
):
    R.extract_stop.clear()
    out_dir = output_dir or DATA_ROOT
    try:
        R.progress["extract"]["running"] = True
        set_progress("extract", 5, "starting", "Preparing extraction")
        log("status", f"Extraction started for {pdf_path.name}")

        if not R.vlm_loaded or R.vlm is None:
            raise RuntimeError("VLM not loaded — click Load VLM first")

        reader = PdfReader(str(pdf_path))
        total = len(reader.pages)

        # Build candidate page list from user range, then subtract already-extracted.
        # Use the PDF's own stem for the lookup so a new/unrecognised file shows
        # all pages as missing rather than falsely reporting them as done.
        _from = max(1, page_from)
        _to = min(total, page_to) if page_to > 0 else total
        pages = list(range(_from, _to + 1))

        # Check by PDF stem in R.docs AND by scanning output_dir on disk
        existing: set[int] = set()
        if stem in R.docs:
            existing |= {p for p, flags in R.docs[stem].items()
                         if "md" in flags or "json" in flags}
        # Also scan output dir for _pageN files matching the stem
        for ext in ("md", "json"):
            for f in out_dir.glob(f"{stem}_page*.{ext}"):
                m = re.match(rf"^{re.escape(stem)}_page(\d+)\.(?:md|json)$",
                              f.name, flags=re.IGNORECASE)
                if m:
                    existing.add(int(m.group(1)))
        pages = [p for p in pages if p not in existing]

        if not pages:
            set_progress("extract", 100, "completed", "All pages in range already extracted")
            log("status", "No missing pages to extract in selected range")
            return

        log("status", f"Extracting {len(pages)} page(s) (range {_from}-{_to}) of {total} total")
        extracted_count = 0
        for i, p in enumerate(pages, start=1):
            if R.extract_stop.is_set():
                set_progress("extract", int((i - 1) / max(1, len(pages)) * 100), "stopped", "Stopped by user")
                log("status", f"Extraction stopped after {extracted_count} page(s)")
                return
            pct = 10 + int((i - 1) / max(1, len(pages)) * 80)
            set_progress("extract", pct, "extracting", f"Page {p}/{total}")
            res = R.vlm.extract_page(str(pdf_path), p, prompt=prompt)
            (out_dir / f"{stem}_page{p}.json").write_text(json.dumps(res, indent=2, ensure_ascii=False), encoding="utf-8")
            (out_dir / f"{stem}_page{p}.md").write_text(str(res.get("raw_response", "")), encoding="utf-8")
            extracted_count += 1

        R.docs = scan_docs(out_dir)
        set_progress("extract", 100, "completed", f"Extracted {extracted_count} page(s)")
        log("status", f"Extraction completed ({extracted_count} page(s))")
    except Exception as e:
        log("error", f"Extraction failed: {e}")
        log("debug", traceback.format_exc())
        set_progress("extract", 100, "failed", str(e))
    finally:
        R.progress["extract"]["running"] = False
        if is_tmp:
            try:
                if pdf_path.exists():
                    pdf_path.unlink()
            except Exception:
                pass


def _resolve_export_dir(output_dir: str) -> Path:
    if output_dir:
        p = Path(output_dir)
        p.mkdir(parents=True, exist_ok=True)
        return p
    return EXPORT_DIR


def export_excel(hits: list[dict[str, Any]], output_dir: str = "") -> Path:
    rows = []
    for h in hits:
        m = h["meta"]
        rows.append(
            {
                "file_name": m.get("file_name"),
                "report_name": m.get("report_name"),
                "page_number": m.get("page_number"),
                "table_id": m.get("table_id"),
                "extraction_type": m.get("extraction_type"),
                "title": m.get("title"),
                "sample_id": m.get("sample_id"),
                "score": h.get("score"),
                "text": h.get("text"),
            }
        )
    out = _resolve_export_dir(output_dir)
    path = out / f"retrieved_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    pd.DataFrame(rows).to_excel(path, index=False)
    return path


def export_word(hits: list[dict[str, Any]], output_dir: str = "") -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("Retrieved SCAL Results", level=1)
    for i, h in enumerate(hits, start=1):
        m = h["meta"]
        doc.add_heading(f"Result {i}", level=2)
        doc.add_paragraph(
            f"File: {m.get('file_name')} | Report: {m.get('report_name')} | "
            f"Page: {m.get('page_number')} | Table: {m.get('table_id')}"
        )
        if m.get("raw_html") and m.get("parsed_rows"):
            rows = m.get("parsed_rows") or []
            cols = m.get("parsed_columns") or (list(rows[0].keys()) if rows else [])
            if cols:
                t = doc.add_table(rows=1, cols=len(cols))
                for ci, c in enumerate(cols):
                    t.rows[0].cells[ci].text = str(c)
                for r in rows:
                    rr = t.add_row().cells
                    for ci, c in enumerate(cols):
                        rr[ci].text = str(r.get(c, ""))
        else:
            doc.add_paragraph(str(h.get("text", ""))[:2000])
    out = _resolve_export_dir(output_dir)
    path = out / f"retrieved_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(path)
    return path


app = FastAPI(title="SCAL Modern Local App")
static_dir = Path(__file__).resolve().parent / "static"
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")


@app.get("/")
def ui_index():
    return FileResponse(static_dir / "index.html")


class ChatReq(BaseModel):
    doc_name: str | None = None
    scope: str = "selected"  # selected | all
    session_id: str | None = None
    question: str
    prompt_template: str = ""
    filter_extraction_type: str | None = None
    top_k: int = 8
    response_mode: str = "balanced"  # fast | balanced | deep


class ExperimentReq(BaseModel):
    mode: str = "full"  # full | retrieval | reasoning
    data_root: str = ""
    benchmark_path: str
    output_root: str = ""
    top_k: int = 3
    run_id: str = ""
    retrieval_configs: list[str] = []
    model_names: list[str] = []
    prompt_types: list[str] = []
    selected_retrieval_config: str = ""


class ModelBenchmarkReq(BaseModel):
    question: str
    simple_context: str = ""
    detailed_context: str = ""
    expected_keywords: str = ""
    model_names: list[str] = []
    run_id: str = ""
    output_root: str = ""


# ── Folder / file browse (non-blocking, runs tkinter in executor) ─────────────

def _tkdialog_folder() -> str:
    import tkinter
    import tkinter.filedialog
    root = tkinter.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    path = tkinter.filedialog.askdirectory(parent=root) or ""
    root.destroy()
    return path


def _tkdialog_file(accept: str) -> str:
    import tkinter
    import tkinter.filedialog
    root = tkinter.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    ftypes = [("PDF files", "*.pdf"), ("All files", "*.*")] if ".pdf" in accept else [("All files", "*.*")]
    path = tkinter.filedialog.askopenfilename(parent=root, filetypes=ftypes) or ""
    root.destroy()
    return path


@app.post("/api/browse/folder")
async def api_browse_folder():
    """Opens a native folder picker dialog (non-blocking)."""
    loop = asyncio.get_event_loop()
    path = await loop.run_in_executor(_DIALOG_EXECUTOR, _tkdialog_folder)
    return {"path": path}


@app.post("/api/browse/file")
async def api_browse_file(payload: dict = {}):
    accept = (payload or {}).get("accept", "")
    loop = asyncio.get_event_loop()
    path = await loop.run_in_executor(_DIALOG_EXECUTOR, _tkdialog_file, accept)
    return {"path": path}


@app.get("/api/settings")
def api_settings_get():
    return {"advanced_mode": R.advanced_mode}


@app.post("/api/settings/advanced-mode")
def api_settings_advanced_mode(enabled: str = Form("false")):
    val = str(enabled).strip().lower() in {"1", "true", "yes", "on"}
    R.advanced_mode = val
    _save_settings({"advanced_mode": val})
    return {"ok": True, "advanced_mode": val}


# ── Documents & index ─────────────────────────────────────────────────────────

@app.get("/api/docs")
async def api_docs(root: str | None = None):
    loop = asyncio.get_event_loop()
    rr = Path(root) if root else DATA_ROOT

    def _scan():
        R.docs = scan_docs(rr)
        names = sorted(R.docs.keys())
        pages_map: dict[str, dict[int, dict[str, str]]] = {
            doc_name: {
                pg: {ext: str(path) for ext, path in files.items()}
                for pg, files in pages.items()
            }
            for doc_name, pages in R.docs.items()
        }
        return names, pages_map

    names, pages_map = await loop.run_in_executor(_SEARCH_EXECUTOR, _scan)
    return {
        "data_root": str(rr),
        "documents": names,
        "coverage": {n: coverage_for_doc(n) for n in names},
        "pages_map": pages_map,
    }


@app.get("/api/docs/debug")
async def api_docs_debug(root: str | None = None):
    """Returns raw file listing to help diagnose why documents aren't appearing."""
    loop = asyncio.get_event_loop()
    rr = Path(root) if root else DATA_ROOT
    if not rr.exists():
        return {"error": f"Path does not exist: {rr}", "path": str(rr)}

    def _list():
        out = []
        for p in rr.iterdir():
            if p.is_file():
                stem, page, ext = parse_name(p.name)
                out.append({"name": p.name, "parsed_stem": stem,
                             "parsed_page": page, "parsed_ext": ext,
                             "matched": stem is not None})
        return out

    files = await loop.run_in_executor(_SEARCH_EXECUTOR, _list)
    return {
        "path": str(rr),
        "total_files": len(files),
        "matched_files": sum(1 for f in files if f["matched"]),
        "files": files[:50],
    }


@app.post("/api/index/build")
def api_build_index(doc_name: str = Form(...)):
    if doc_name not in R.docs:
        raise HTTPException(status_code=404, detail="Document not found")
    if R.progress["index"]["running"]:
        return {"ok": False, "message": "Index build already running"}
    t = threading.Thread(target=build_index_job, args=(doc_name,), daemon=True)
    t.start()
    return {"ok": True, "message": f"Index build started for {doc_name}"}


def build_all_index_job(data_root: Path):
    """Index ALL documents in the folder sequentially."""
    docs = scan_docs(data_root)
    names = sorted(docs.keys())
    if not names:
        log("error", "No documents found in folder")
        return
    log("status", f"Building index for all {len(names)} document(s)")
    R.docs = docs
    total_chunks = 0
    for i, doc_name in enumerate(names):
        if R.progress["index"]["running"]:
            # Wait briefly if another build is still finishing
            import time; time.sleep(0.2)
        R.progress["index"]["running"] = True
        set_progress("index", int(i / len(names) * 90), "indexing", f"({i+1}/{len(names)}) {doc_name}")
        try:
            chunks = chunks_for_doc(doc_name)
            if not chunks:
                continue
            texts = [c["text"] for c in chunks]
            metas = [c["meta"] for c in chunks]
            vec = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
            mat = vec.fit_transform(texts)
            save_index(ns(doc_name), vec, mat, texts, metas)
            total_chunks += len(chunks)
            log("status", f"Indexed {doc_name}: {len(chunks)} chunks")
        except Exception as e:
            log("error", f"Index failed for {doc_name}: {e}")
        finally:
            R.progress["index"]["running"] = False
    try:
        set_progress("index", 95, "indexing", "Building global all-doc index")
        total_chunks = build_global_index_job(names)
    except Exception as e:
        log("error", f"Global all-doc index build failed: {e}")
    set_progress("index", 100, "completed", f"All docs indexed — {total_chunks} total chunks")
    log("status", f"All-docs index complete: {total_chunks} chunks")


@app.post("/api/index/build-all")
def api_build_all(data_root: str = Form("")):
    if R.progress["index"]["running"]:
        return {"ok": False, "message": "Index build already running"}
    rr = Path(data_root) if data_root else DATA_ROOT
    t = threading.Thread(target=build_all_index_job, args=(rr,), daemon=True)
    t.start()
    return {"ok": True, "message": f"All-docs index build started for {rr}"}


# ── State & logs ──────────────────────────────────────────────────────────────

@app.get("/api/state")
def api_state():
    sync_llm_state_from_inference()
    return {
        "progress": R.progress,
        "advanced_mode": R.advanced_mode,
        "experiment": R.experiment,
        "benchmark": R.benchmark,
        "models": {
            "llm_loaded": R.llm_loaded,
            "llm_model": R.llm_model_name,
            "llm_target_model": R.llm_target_model,
            "llm_last_error": R.llm_last_error,
            "llm_loading": bool(R.progress.get("model", {}).get("running", False)),
            "vlm_loaded": R.vlm_loaded,
        },
    }


@app.get("/api/logs")
def api_logs(kind: str = "status", limit: int = 200):
    if kind not in R.logs:
        raise HTTPException(status_code=400, detail="Invalid log kind")
    items = list(R.logs[kind])[-limit:]
    return {"kind": kind, "items": items}


@app.post("/api/logs/clear")
def api_clear_logs(kind: str = Form("all")):
    if kind == "all":
        for k in R.logs:
            R.logs[k].clear()
    elif kind in R.logs:
        R.logs[kind].clear()
    else:
        raise HTTPException(status_code=400, detail="Invalid log kind")
    return {"ok": True}


# ── Models ────────────────────────────────────────────────────────────────────

def _load_llm_bg(model_name: str):
    """Run in a daemon thread so the HTTP response returns immediately."""
    try:
        load_llm(model_name)
    except Exception as e:
        msg = str(e)
        low = msg.lower()
        hint = ""
        if "outofmemory" in low or "cuda out of memory" in low:
            hint = "GPU OOM: choose a smaller model or unload and retry"
        elif "no module named" in low:
            hint = "Missing dependency in current environment"
        elif "no space left" in low or "disk" in low:
            hint = "Insufficient disk space for model shards"
        elif "401" in low or "403" in low or "gated" in low:
            hint = "Model access/auth issue on Hugging Face"
        elif "timeout" in low or "connection" in low:
            hint = "Network timeout during model download"
        final_err = f"{msg}" if not hint else f"{msg} ({hint})"
        R.llm_loaded = False
        R.llm_last_error = final_err
        R.llm_target_model = ""
        R.progress["model"]["running"] = False
        set_progress("model", 100, "failed", f"LLM load failed: {final_err}")
        log("error", f"LLM load failed: {final_err}")
        log("debug", traceback.format_exc())


def _load_vlm_bg():
    try:
        load_vlm()
    except Exception as e:
        log("error", f"VLM load failed: {e}")
        log("debug", traceback.format_exc())


@app.post("/api/models/load-llm")
def api_load_llm(model_name: str = Form("Qwen/Qwen2.5-14B-Instruct")):
    model_name = (model_name or "Qwen/Qwen2.5-14B-Instruct").strip()
    sync_llm_state_from_inference()
    if R.llm_loaded and R.llm_model_name == model_name:
        return {"ok": True, "message": "LLM already loaded"}
    try:
        resp = load_llm(model_name)
        msg = str(resp.get("message") or f"LLM loading started: {model_name}")
        log("status", msg)
        return {"ok": True, "message": msg}
    except Exception as e:
        R.llm_last_error = str(e)
        log("error", f"LLM load request failed: {e}")
        return {"ok": False, "message": str(e)}


@app.post("/api/models/switch-llm")
def api_switch_llm(model_name: str = Form("Qwen/Qwen2.5-14B-Instruct")):
    model_name = (model_name or "Qwen/Qwen2.5-14B-Instruct").strip()
    sync_llm_state_from_inference()
    if R.llm_loaded and R.llm_model_name == model_name:
        return {"ok": True, "message": "LLM already active"}
    try:
        resp = load_llm(model_name)
        msg = str(resp.get("message") or f"Switching LLM to {model_name}")
        log("status", msg)
        return {"ok": True, "message": msg}
    except Exception as e:
        R.llm_last_error = str(e)
        log("error", f"LLM switch request failed: {e}")
        return {"ok": False, "message": str(e)}


@app.post("/api/models/unload-llm")
def api_unload_llm():
    sync_llm_state_from_inference()
    if not R.llm_loaded:
        return {"ok": True, "message": "LLM already unloaded"}
    try:
        resp = unload_llm()
        msg = str(resp.get("message") or "LLM unloaded")
        log("status", msg)
        return {"ok": True, "message": msg}
    except Exception as e:
        log("error", f"LLM unload failed: {e}")
        return {"ok": False, "message": f"LLM unload failed: {e}"}


@app.get("/api/models/options")
def api_model_options():
    try:
        resp = _inference_request("GET", "/v1/models", timeout=8)
        return {
            "models": resp.get("models", LLM_MODEL_OPTIONS),
            "default": resp.get("default", R.llm_model_name),
            "active": resp.get("active", ""),
        }
    except Exception:
        return {"models": LLM_MODEL_OPTIONS, "default": R.llm_model_name, "active": ""}


@app.post("/api/models/load-vlm")
def api_load_vlm():
    if R.vlm_loaded:
        return {"ok": True, "message": "VLM already loaded"}
    log("status", "VLM load requested")
    threading.Thread(target=_load_vlm_bg, daemon=True).start()
    return {"ok": True, "message": "VLM loading in background — watch logs / model pill"}


@app.get("/api/experiments/state")
def api_experiment_state():
    return R.experiment


@app.get("/api/experiments/options")
def api_experiment_options():
    models = [m for m in LLM_MODEL_OPTIONS if not m.get("name", "").startswith("moonshotai/Kimi")]
    return {
        "retrieval_configs": RETRIEVAL_CONFIGS,
        "prompt_types": PROMPT_TYPES,
        "models": models,
    }


@app.post("/api/experiments/stop")
def api_experiment_stop():
    if not R.experiment.get("running"):
        return {"ok": True, "message": "No experiment running"}
    R.experiment_stop.set()
    return {"ok": True, "message": "Experiment stop requested"}


@app.post("/api/experiments/run")
def api_experiment_run(req: ExperimentReq):
    if not R.advanced_mode:
        raise HTTPException(status_code=403, detail="Advanced mode is disabled")
    if R.experiment.get("running"):
        return {"ok": False, "message": "Experiment already running"}
    mode = (req.mode or "full").strip().lower()
    if mode not in {"full", "retrieval", "reasoning"}:
        raise HTTPException(status_code=400, detail="Invalid mode")
    payload = req.model_dump()
    t = threading.Thread(target=_experiment_worker, args=(payload,), daemon=True)
    t.start()
    return {"ok": True, "message": f"Experiment started ({mode})", "run_id": req.run_id or "auto"}


@app.get("/api/benchmarks/models/state")
def api_model_benchmark_state():
    return R.benchmark


@app.post("/api/benchmarks/models/stop")
def api_model_benchmark_stop():
    if not R.benchmark.get("running"):
        return {"ok": True, "message": "No benchmark running"}
    R.benchmark_stop.set()
    return {"ok": True, "message": "Benchmark stop requested"}


@app.post("/api/benchmarks/models/run")
def api_model_benchmark_run(req: ModelBenchmarkReq):
    sync_llm_state_from_inference()
    if R.progress.get("model", {}).get("running") or R.benchmark.get("running"):
        return {"ok": False, "message": "Model operation already running"}
    payload = req.model_dump()
    t = threading.Thread(target=_run_model_benchmark_worker, args=(payload,), daemon=True)
    t.start()
    return {"ok": True, "message": "Model benchmark started", "run_id": req.run_id or "auto"}


# ── Chat ──────────────────────────────────────────────────────────────────────

@app.post("/api/chat")
async def api_chat(req: ChatReq):
    """
    Fully async chat handler:
    - RAG search runs in _SEARCH_EXECUTOR (CPU, non-blocking to event loop)
    - LLM inference runs in _INFERENCE_EXECUTOR (GPU, serialised, non-blocking)
    Both executors are independent of the extraction daemon thread, so PDF
    extraction and chat can run simultaneously without interfering.
    """
    t_total0 = time.perf_counter()
    retrieval_ms = 0.0
    generation_ms = 0.0
    loop = asyncio.get_event_loop()
    filters = {"extraction_type": req.filter_extraction_type}
    scope = (req.scope or "selected").lower()
    response_mode = (req.response_mode or "balanced").lower()
    gen_cfg = llm_generation_settings(response_mode)
    target_doc = "__ALL__" if scope == "all" else (req.doc_name or "")

    # If no selected doc is provided, gracefully fall back to all-doc scope
    # so normal chat still works without forcing user selection.
    if scope != "all" and not target_doc:
        scope = "all"
        target_doc = "__ALL__"

    # Lazy doc refresh so chat can still work after restart
    if target_doc == "__ALL__" and not R.docs:
        try:
            R.docs = scan_docs(DATA_ROOT)
        except Exception:
            R.docs = {}

    # Casual chat mode: avoid retrieving random PDF chunks for greetings/small talk
    if is_casual_chat(req.question):
        try:
            loop = asyncio.get_event_loop()
            t_g0 = time.perf_counter()
            answer = await loop.run_in_executor(
                _INFERENCE_EXECUTOR,
                lambda: ask_llm(
                    "You are a friendly assistant in a SCAL document app. "
                    "For casual chat, respond naturally and briefly. "
                    "Do not cite PDFs unless user asks document questions.",
                    req.question,
                    **gen_cfg,
                ),
            )
            generation_ms = (time.perf_counter() - t_g0) * 1000.0
        except Exception:
            answer = (
                "Hi! I can chat casually, and when you're ready I can also help query your extracted SCAL PDFs. "
                "Try asking about porosity, permeability, capillary pressure, or specific samples."
            )

        sid = req.session_id
        if not sid:
            s = create_session("SCAL Chat Session")
            sid = s["id"]
        append_session_messages(
            sid,
            [
                {"role": "user", "content": req.question, "time": now()},
                {"role": "assistant", "content": answer, "time": now(), "sources": []},
            ],
        )

        return {
            "session_id": sid,
            "answer": answer,
            "reasoning": [],
            "sources": [],
            "tables": [],
            "raw_hits": [],
            "metrics": {
                "response_mode": response_mode,
                "model_name": R.llm_model_name,
                "retrieval_ms": round(retrieval_ms, 2),
                "generation_ms": round(generation_ms, 2),
                "total_ms": round((time.perf_counter() - t_total0) * 1000.0, 2),
                "answer_tokens": approx_token_count(answer),
                "tokens_per_sec": round((approx_token_count(answer) / max(generation_ms / 1000.0, 1e-6)), 2)
                if generation_ms > 0
                else 0.0,
                "hits": 0,
            },
        }

    # ── 1. RAG search (CPU-bound TF-IDF) ──────────────────────────────────────
    t_r0 = time.perf_counter()
    hits: list[dict[str, Any]] = await loop.run_in_executor(
        _SEARCH_EXECUTOR,
        lambda: search(req.question, target_doc, filters, top_k=req.top_k),
    )
    retrieval_ms = (time.perf_counter() - t_r0) * 1000.0

    # ── 2. Build reasoning list ────────────────────────────────────────────────
    reasoning = []
    for i, h in enumerate(hits, start=1):
        m = h["meta"]
        reasoning.append(
            {
                "rank": i,
                "score": round(h["score"], 4),
                "file_name": m.get("file_name"),
                "page_number": m.get("page_number"),
                "table_id": m.get("table_id"),
                "extraction_type": m.get("extraction_type"),
                "snippet": str(h.get("text", ""))[:300],
            }
        )

    # ── 3. LLM answer (GPU-bound, serialised) ─────────────────────────────────
    if not hits:
        # No retrieval evidence: fall back to normal chat answer instead of blocking.
        # This keeps chat usable when no docs are loaded/indexed.
        system = (
            "You are a helpful assistant in a SCAL app. "
            "If no document evidence is available, answer generally and clearly state "
            "that the response is not grounded in retrieved files."
        )
        user_prompt = req.question
        try:
            t_g0 = time.perf_counter()
            answer = await loop.run_in_executor(
                _INFERENCE_EXECUTOR,
                lambda: ask_llm(system, user_prompt, **gen_cfg),
            )
            generation_ms = (time.perf_counter() - t_g0) * 1000.0
            answer = (
                "(No retrieved document context found; general model response)\n\n"
                + answer
            )
        except Exception as e:
            answer = f"No retrieved chunks, and model fallback failed: {e}"
    else:
        ctx = []
        for i, h in enumerate(hits, start=1):
            m = h["meta"]
            txt = h["text"]
            if len(txt) > 700:
                txt = txt[:700] + "..."
            ctx.append(
                f"[{i}] file={m.get('file_name')} page={m.get('page_number')} table={m.get('table_id')}\n{txt}"
            )
        context = "\n\n".join(ctx)
        system = "You are a SCAL assistant. Use only retrieved evidence and cite [1],[2]."
        user_prompt = f"Task prompt:\n{req.prompt_template}\n\nQuestion:\n{req.question}\n\nEvidence:\n{context}"
        try:
            t_g0 = time.perf_counter()
            answer = await loop.run_in_executor(
                _INFERENCE_EXECUTOR,
                lambda: ask_llm(system, user_prompt, **gen_cfg),
            )
            generation_ms = (time.perf_counter() - t_g0) * 1000.0
        except Exception as e:
            answer = f"LLM not loaded or failed: {e}\n\nFallback evidence:\n{context[:1800]}"

    # ── 4. Collect tables from hits ────────────────────────────────────────────
    tables = []
    for h in hits:
        m = h["meta"]
        if m.get("parsed_rows") or m.get("raw_html"):
            tables.append(
                {
                    "file_name": m.get("file_name"),
                    "report_name": m.get("report_name"),
                    "page_number": m.get("page_number"),
                    "table_id": m.get("table_id"),
                    "raw_html": m.get("raw_html"),
                    "columns": m.get("parsed_columns"),
                    "rows": m.get("parsed_rows"),
                }
            )

    # Persist chat memory (session-based)
    sid = req.session_id
    if not sid:
        s = create_session("SCAL Chat Session")
        sid = s["id"]
    append_session_messages(
        sid,
        [
            {"role": "user", "content": req.question, "time": now()},
            {"role": "assistant", "content": answer, "time": now(), "sources": reasoning},
        ],
    )

    answer_tokens = approx_token_count(answer)
    total_ms = (time.perf_counter() - t_total0) * 1000.0

    return {
        "session_id": sid,
        "answer": answer,
        "reasoning": reasoning,
        "sources": reasoning,
        "tables": tables,
        "raw_hits": hits,
        "metrics": {
            "response_mode": response_mode,
            "model_name": R.llm_model_name,
            "retrieval_ms": round(retrieval_ms, 2),
            "generation_ms": round(generation_ms, 2),
            "total_ms": round(total_ms, 2),
            "answer_tokens": int(answer_tokens),
            "tokens_per_sec": round((answer_tokens / max(generation_ms / 1000.0, 1e-6)), 2)
            if generation_ms > 0
            else 0.0,
            "hits": len(hits),
            "max_new_tokens": int(gen_cfg.get("max_new_tokens", 0)),
        },
    }


@app.post("/api/chat/stream")
async def api_chat_stream(req: ChatReq):
    t_total0 = time.perf_counter()
    retrieval_ms = 0.0
    loop = asyncio.get_event_loop()
    filters = {"extraction_type": req.filter_extraction_type}
    scope = (req.scope or "selected").lower()
    response_mode = (req.response_mode or "balanced").lower()
    gen_cfg = llm_generation_settings(response_mode)
    target_doc = "__ALL__" if scope == "all" else (req.doc_name or "")
    if scope != "all" and not target_doc:
        target_doc = "__ALL__"
    if target_doc == "__ALL__" and not R.docs:
        try:
            R.docs = scan_docs(DATA_ROOT)
        except Exception:
            R.docs = {}

    t_r0 = time.perf_counter()
    hits: list[dict[str, Any]] = []
    if not is_casual_chat(req.question):
        hits = await loop.run_in_executor(
            _SEARCH_EXECUTOR,
            lambda: search(req.question, target_doc, filters, top_k=req.top_k),
        )
    retrieval_ms = (time.perf_counter() - t_r0) * 1000.0

    reasoning = []
    for i, h in enumerate(hits, start=1):
        m = h["meta"]
        reasoning.append(
            {
                "rank": i,
                "score": round(h["score"], 4),
                "file_name": m.get("file_name"),
                "page_number": m.get("page_number"),
                "table_id": m.get("table_id"),
                "extraction_type": m.get("extraction_type"),
                "snippet": str(h.get("text", ""))[:300],
            }
        )

    tables = []
    for h in hits:
        m = h["meta"]
        if m.get("parsed_rows") or m.get("raw_html"):
            tables.append(
                {
                    "file_name": m.get("file_name"),
                    "report_name": m.get("report_name"),
                    "page_number": m.get("page_number"),
                    "table_id": m.get("table_id"),
                    "raw_html": m.get("raw_html"),
                    "columns": m.get("parsed_columns"),
                    "rows": m.get("parsed_rows"),
                }
            )

    if is_casual_chat(req.question):
        system = (
            "You are a friendly assistant in a SCAL document app. "
            "For casual chat, respond naturally and briefly. "
            "Do not cite PDFs unless user asks document questions."
        )
        user_prompt = req.question
        no_context_prefix = ""
    elif not hits:
        system = (
            "You are a helpful assistant in a SCAL app. "
            "If no document evidence is available, answer generally and clearly state "
            "that the response is not grounded in retrieved files."
        )
        user_prompt = req.question
        no_context_prefix = "(No retrieved document context found; general model response)\n\n"
    else:
        ctx = []
        for i, h in enumerate(hits, start=1):
            m = h["meta"]
            txt = h["text"]
            if len(txt) > 700:
                txt = txt[:700] + "..."
            ctx.append(
                f"[{i}] file={m.get('file_name')} page={m.get('page_number')} table={m.get('table_id')}\n{txt}"
            )
        context = "\n\n".join(ctx)
        system = "You are a SCAL assistant. Use only retrieved evidence and cite [1],[2]."
        user_prompt = f"Task prompt:\n{req.prompt_template}\n\nQuestion:\n{req.question}\n\nEvidence:\n{context}"
        no_context_prefix = ""

    sid = req.session_id
    if not sid:
        s = create_session("SCAL Chat Session")
        sid = s["id"]

    payload = {
        "system_prompt": system,
        "user_prompt": user_prompt,
        "max_new_tokens": int(gen_cfg.get("max_new_tokens", 420)),
        "temperature": float(gen_cfg.get("temperature", 0.2)),
        "top_p": float(gen_cfg.get("top_p", 0.9)),
        "do_sample": bool(gen_cfg.get("do_sample", True)),
    }

    def _sse(evt: dict[str, Any]) -> str:
        return f"data: {json.dumps(evt, ensure_ascii=False)}\n\n"

    def streamer():
        answer_parts: list[str] = []
        infer_metrics: dict[str, Any] = {}
        try:
            if no_context_prefix:
                answer_parts.append(no_context_prefix)
                yield _sse({"type": "token", "text": no_context_prefix})

            for ev in _inference_stream_events("/v1/chat/stream", payload, timeout=600):
                if ev.get("type") == "token":
                    piece = str(ev.get("text") or "")
                    if piece:
                        answer_parts.append(piece)
                        yield _sse({"type": "token", "text": piece})
                elif ev.get("type") == "metrics":
                    infer_metrics = ev.get("metrics") or {}

            answer = "".join(answer_parts).strip()
            append_session_messages(
                sid,
                [
                    {"role": "user", "content": req.question, "time": now()},
                    {"role": "assistant", "content": answer, "time": now(), "sources": reasoning},
                ],
            )

            generation_ms = float(infer_metrics.get("generation_ms", 0) or 0)
            answer_tokens = int(infer_metrics.get("answer_tokens", approx_token_count(answer)) or 0)
            tok_s = float(infer_metrics.get("tokens_per_sec", 0) or 0)
            metrics = {
                "response_mode": response_mode,
                "model_name": R.llm_model_name,
                "retrieval_ms": round(retrieval_ms, 2),
                "generation_ms": round(generation_ms, 2),
                "first_token_ms": round(float(infer_metrics.get("first_token_ms", generation_ms) or generation_ms), 2),
                "total_ms": round((time.perf_counter() - t_total0) * 1000.0, 2),
                "answer_tokens": answer_tokens,
                "tokens_per_sec": round(tok_s, 2) if tok_s else round((answer_tokens / max(generation_ms / 1000.0, 1e-6)), 2),
                "hits": len(hits),
                "max_new_tokens": int(gen_cfg.get("max_new_tokens", 0)),
            }
            done = {
                "type": "done",
                "session_id": sid,
                "answer": answer,
                "reasoning": reasoning,
                "sources": reasoning,
                "tables": tables,
                "raw_hits": hits,
                "metrics": metrics,
            }
            yield _sse(done)
        except Exception as e:
            yield _sse({"type": "error", "message": str(e)})

    return StreamingResponse(streamer(), media_type="text/event-stream")


@app.get("/api/chat/sessions")
def api_chat_sessions():
    return {"sessions": list_sessions()}


@app.post("/api/chat/session/new")
def api_chat_session_new(title: str = Form("")):
    s = create_session(title)
    return {"session": {"id": s["id"], "title": s["title"], "messages": s["messages"]}}


@app.get("/api/chat/session/{session_id}")
def api_chat_session_get(session_id: str):
    s = get_session(session_id)
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")
    return {"session": s}


@app.get("/api/chat/suggestions")
def api_chat_suggestions(doc_name: str | None = None):
    # Static use-case prompts (user-facing labels + hidden prompt templates)
    out = list(USE_CASE_PROMPT_SUGGESTIONS)

    # Dynamic suggestions from current RAG metadata
    dynamic = []
    docs = [doc_name] if doc_name and doc_name in R.docs else list(R.docs.keys())[:20]
    type_counts: dict[str, int] = {}
    for d in docs:
        pages = R.docs.get(d, {})
        for pg in pages.values():
            md = pg.get("md")
            js = pg.get("json")
            txt = ""
            if md and md.exists():
                txt = md.read_text(encoding="utf-8", errors="ignore")
            elif js and js.exists():
                try:
                    txt = flatten_json(json.loads(js.read_text(encoding="utf-8", errors="ignore")))
                except Exception:
                    txt = js.read_text(encoding="utf-8", errors="ignore")
            t = infer_type(txt)
            type_counts[t] = type_counts.get(t, 0) + 1
    for t, n in sorted(type_counts.items(), key=lambda x: x[1], reverse=True)[:4]:
        dynamic.append(
            {
                "id": f"dyn_{t}",
                "label": f"Explore {t.replace('_', ' ')} ({n} pages)",
                "question": f"Summarize key findings for {t.replace('_', ' ')} with sources.",
                "prompt_template": "Return concise technical bullets, include key values and cite sources [1],[2].",
            }
        )
    return {"suggestions": out + dynamic}


# ── Export ────────────────────────────────────────────────────────────────────

@app.post("/api/export/excel")
def api_export_excel(payload: dict):
    path = export_excel(payload.get("hits", []), output_dir=payload.get("output_dir", ""))
    return {"ok": True, "path": str(path)}


@app.post("/api/export/word")
def api_export_word(payload: dict):
    path = export_word(payload.get("hits", []), output_dir=payload.get("output_dir", ""))
    return {"ok": True, "path": str(path)}


# ── PDF Extraction ────────────────────────────────────────────────────────────

def _pdf_check(pdf_path: Path, doc_name: str, output_dir: str = "",
               stem_override: str = "") -> dict:
    reader = PdfReader(str(pdf_path))
    total = len(reader.pages)

    # Use the real filename stem (may differ from the temp path stem for uploads)
    pdf_stem = stem_override or pdf_path.stem

    # Search both the configured output_dir and DATA_ROOT for already-extracted pages
    search_roots = [DATA_ROOT]
    if output_dir:
        out = Path(output_dir)
        if out.exists() and out != DATA_ROOT:
            search_roots.insert(0, out)

    extracted: set[int] = set()

    # 1. Check R.docs by PDF stem (most reliable when already scanned)
    if pdf_stem in R.docs:
        extracted |= {p for p, flags in R.docs[pdf_stem].items()
                      if "md" in flags or "json" in flags}

    # 2. Also scan roots directly for _pageN.md / _pageN.json files matching the stem
    for root in search_roots:
        if not root.exists():
            continue
        for ext in ("md", "json"):
            for f in root.glob(f"{pdf_stem}_page*.{ext}"):
                m = re.match(rf"^{re.escape(pdf_stem)}_page(\d+)\.(?:md|json)$",
                              f.name, flags=re.IGNORECASE)
                if m:
                    extracted.add(int(m.group(1)))

    missing = [p for p in range(1, total + 1) if p not in extracted]
    return {
        "total_pdf_pages": total,
        "already_extracted": len(missing) == 0,
        "missing_pages": missing,
        "extracted_pages": sorted(extracted),
        "pdf_stem": pdf_stem,
    }


@app.post("/api/extract/check")
async def api_extract_check(
    file: UploadFile = File(None),
    pdf_path: str = Form(""),
    doc_name: str = Form(""),
    output_dir: str = Form(""),
):
    if file and file.filename:
        real_stem = Path(file.filename).stem
        data = await file.read()
        with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        try:
            return _pdf_check(tmp_path, doc_name, output_dir, stem_override=real_stem)
        finally:
            try: tmp_path.unlink()
            except Exception: pass
    elif pdf_path:
        p = Path(pdf_path)
        if not p.exists():
            raise HTTPException(status_code=400, detail=f"File not found: {pdf_path}")
        return _pdf_check(p, doc_name, output_dir)
    else:
        raise HTTPException(status_code=400, detail="Provide file upload or pdf_path")


@app.post("/api/extract/start")
async def api_extract_start(
    file: UploadFile = File(None),
    pdf_path: str = Form(""),
    prompt: str = Form(""),
    target_doc_name: str = Form(""),
    output_dir: str = Form(""),
    page_from: int = Form(1),
    page_to: int = Form(0),
):
    if R.progress["extract"]["running"]:
        return JSONResponse({"ok": False, "message": "Extraction already running"}, status_code=409)

    if file and file.filename:
        data = await file.read()
        stem = Path(file.filename).stem
        with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        is_tmp = True
    elif pdf_path:
        p = Path(pdf_path)
        if not p.exists():
            raise HTTPException(status_code=400, detail=f"File not found: {pdf_path}")
        stem = p.stem
        tmp_path = p
        is_tmp = False
    else:
        raise HTTPException(status_code=400, detail="Provide file upload or pdf_path")

    out = Path(output_dir) if output_dir else None
    t = threading.Thread(
        target=extraction_job,
        args=(tmp_path, stem, prompt, target_doc_name or None, out, page_from, page_to, is_tmp),
        daemon=True,
    )
    t.start()
    return {"ok": True, "message": f"Extraction started ({stem})"}


@app.post("/api/extract/stop")
def api_extract_stop():
    R.extract_stop.set()
    return {"ok": True, "message": "Stop signal sent"}


# ── Viewer endpoints ──────────────────────────────────────────────────────────

@app.get("/api/page/raw")
async def api_page_raw(doc: str, page: int):
    """Return the raw text content (JSON or MD) for a specific page."""
    pages = R.docs.get(doc, {})
    if page not in pages:
        raise HTTPException(status_code=404, detail="Page not found")
    files = pages[page]

    def _read():
        if "json" in files:
            content = files["json"].read_text(encoding="utf-8", errors="ignore")
            try:
                content = json.dumps(json.loads(content), indent=2, ensure_ascii=False)
            except Exception:
                pass
            return {"kind": "json", "content": content}
        elif "md" in files:
            content = files["md"].read_text(encoding="utf-8", errors="ignore")
            return {"kind": "md", "content": content}
        return {"kind": "none", "content": "(no extracted file for this page)"}

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(_SEARCH_EXECUTOR, _read)


@app.get("/api/page/pdf")
def api_page_pdf(doc: str, page: int):
    """Serve the raw page PDF file for preview (FileResponse is already async-friendly)."""
    pages = R.docs.get(doc, {})
    if page not in pages:
        raise HTTPException(status_code=404, detail="Page not found")
    files = pages[page]
    if "pdf" not in files:
        raise HTTPException(status_code=404, detail="No PDF for this page")
    return FileResponse(str(files["pdf"]), media_type="application/pdf")


@app.get("/api/page/parse")
async def api_page_parse(doc: str, page: int):
    """Parse HTML tables from the extracted content of a page."""
    pages = R.docs.get(doc, {})
    if page not in pages:
        raise HTTPException(status_code=404, detail="Page not found")
    files = pages[page]

    def _parse():
        raw = ""
        if "json" in files:
            try:
                raw = flatten_json(json.loads(files["json"].read_text(encoding="utf-8", errors="ignore")))
            except Exception:
                raw = files["json"].read_text(encoding="utf-8", errors="ignore")
        elif "md" in files:
            raw = files["md"].read_text(encoding="utf-8", errors="ignore")
        html_tables = extract_html_tables(raw)
        result = []
        for h in html_tables:
            cols, rows = parse_html_table(h)
            result.append({"raw_html": h, "columns": cols, "rows": rows, "row_count": len(rows)})
        return result

    loop = asyncio.get_event_loop()
    tables = await loop.run_in_executor(_SEARCH_EXECUTOR, _parse)
    return {"page": page, "tables": tables, "table_count": len(tables)}
