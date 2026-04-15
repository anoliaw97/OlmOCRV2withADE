from __future__ import annotations

from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path

import pandas as pd


@dataclass(slots=True)
class ChatRecord:
    timestamp: str
    mode: str
    runtime: str
    model: str
    question: str
    answer: str
    citations: str


class ExportService:
    def export_chat_records(self, records: list[ChatRecord], destination: Path) -> tuple[bool, str]:
        if not records:
            return False, "No chat history available to export."

        destination = destination.expanduser().resolve()
        destination.parent.mkdir(parents=True, exist_ok=True)

        if destination.suffix.lower() == ".csv":
            return self._to_csv(records, destination)

        if destination.suffix.lower() in {".xlsx", ".xls"}:
            return self._to_excel(records, destination)

        if destination.suffix.lower() == ".docx":
            return self._to_docx(records, destination)

        return False, f"Unsupported export format: {destination.suffix}"

    def _to_csv(self, records: list[ChatRecord], destination: Path) -> tuple[bool, str]:
        frame = pd.DataFrame([asdict(r) for r in records])
        frame.to_csv(destination, index=False)
        return True, f"Exported CSV: {destination}"

    def _to_excel(self, records: list[ChatRecord], destination: Path) -> tuple[bool, str]:
        frame = pd.DataFrame([asdict(r) for r in records])
        meta = pd.DataFrame(
            [
                {
                    "generated_at": datetime.now().isoformat(timespec="seconds"),
                    "rows": len(frame),
                    "modes": ", ".join(sorted({str(v) for v in frame["mode"].dropna().unique().tolist()})),
                    "models": ", ".join(sorted({str(v) for v in frame["model"].dropna().unique().tolist()})),
                }
            ]
        )

        with pd.ExcelWriter(destination, engine="openpyxl") as writer:
            meta.to_excel(writer, sheet_name="Metadata", index=False)
            frame.to_excel(writer, sheet_name="ChatRecords", index=False)

            for sheet in writer.sheets.values():
                for col in sheet.columns:
                    max_len = 0
                    letter = col[0].column_letter
                    for cell in col[:300]:
                        text = "" if cell.value is None else str(cell.value)
                        if len(text) > max_len:
                            max_len = len(text)
                    sheet.column_dimensions[letter].width = min(70, max(12, max_len + 2))
                sheet.freeze_panes = "A2"
        return True, f"Exported Excel: {destination}"

    def _to_docx(self, records: list[ChatRecord], destination: Path) -> tuple[bool, str]:
        try:
            from docx import Document
        except ImportError:
            return False, "python-docx is not installed. Install requirements again and retry."

        document = Document()
        document.add_heading("Chat Export", level=1)

        for idx, record in enumerate(records, start=1):
            document.add_heading(f"Turn {idx}", level=2)
            document.add_paragraph(f"Timestamp: {record.timestamp}")
            document.add_paragraph(f"Mode: {record.mode}")
            document.add_paragraph(f"Runtime: {record.runtime}")
            document.add_paragraph(f"Model: {record.model}")
            document.add_paragraph(f"Question: {record.question}")
            document.add_paragraph("Answer:")
            document.add_paragraph(record.answer)
            if record.citations.strip():
                document.add_paragraph(f"Citations: {record.citations}")

        document.save(destination)
        return True, f"Exported Word document: {destination}"
